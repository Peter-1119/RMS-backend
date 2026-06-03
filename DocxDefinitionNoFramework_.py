import json, re, os, io, base64
import hashlib
from PIL import Image
from datetime import datetime
from enum import Enum

from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from modules.block_tree import format_node_number  # 階層編號（spec §5）

COLOR_DICT = {
    "red": RGBColor(255, 0, 0),
    "blue": RGBColor(0, 0, 255),
    'black': RGBColor(0, 0, 0),
    "#000": RGBColor(0, 0, 0),
    "#000000": RGBColor(0, 0, 0),
    "#ff0000": RGBColor(255, 0, 0),
    "#0000ff": RGBColor(0, 0, 255),
}

class DOCUMENT_TYPE(Enum):
    ManufacturingDocument = 0
    SpecificationDocument = 1

class DOCUMENT_STEP(Enum):
    ManufacturingDocument = [
        {"目的": {"parent": "attribute", "code": "purpose"}},
        {"製造流程": {"parent": "content", "code": 0}},
        {"管理條件": {"parent": "content", "code": 1}},
        {"製造條件參數一覽表": {"parent": "content", "code": 2}},
        {"異常處置": {"parent": "content", "code": 3}},
        {"相關文件": {"parent": "reference", "code": 0}},
        {"使用表單": {"parent": "reference", "code": 1}}
    ]
    SpecificationDocument = [
        {"目的": {"parent": "attribute", "code": "documentPurpose"}},
        {"製作條件規範": {"parent": "content", "code": 4}},
        {"製造參數一覽表": {"parent": "content", "code": 5}},
        {"適用品質與規格內容": {"parent": "content", "code": 6}},
        {"使用表單": {"parent": "reference", "code": 1}},
        {"其他": {"parent": "content", "code": 7}},
    ]


SPEC_HDRS = ['規格下限(OOS-)','操作下限(OOC-)','設定值','操作上限(OOC+)','規格上限(OOS+)']
COLOR_RED = RGBColor(255, 0, 0)
COLOR_BLUE = RGBColor(0, 0, 255)

def clear_paragraph(p):
    """把一個 paragraph 內原本所有 run/文字清掉"""
    p.text = ""
    # 保險一點也可以把底層 element 都清掉
    for r in p.runs:
        r._element.getparent().remove(r._element)

def add_simple_field(paragraph, instr, font_size=None, hidden=False, color=None):
    """
    在 paragraph 中插入一個簡單欄位，例如:
      instr = "PAGE" 或 "NUMPAGES"
    font_size: Pt(...) / 數字 (例如 10.5)，會轉成半點數給 w:sz
    hidden: True => 使用 w:vanish 隱藏這個 run
    color: RGBColor(...)，可選 (例如 RGBColor(255, 255, 255) = 白色)
    """
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), instr)  # e.g. "PAGE", "NUMPAGES"

    r = OxmlElement('w:r')

    # ---- run properties (rPr) ----
    rPr = OxmlElement('w:rPr')

    # 字型大小
    if font_size is not None:
        # Word 的 w:sz 是「半點數」：10.5pt => 21
        val = int(round(float(font_size) * 2))
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(val))
        rPr.append(sz)

        szCs = OxmlElement('w:szCs')  # complex script size，順便設
        szCs.set(qn('w:val'), str(val))
        rPr.append(szCs)

    # 顏色
    if color is not None:
        color_el = OxmlElement('w:color')
        color_el.set(qn('w:val'), f"{color.rgb:06X}" if hasattr(color, "rgb") else color)
        # 如果用 RGBColor，就使用 color.rgb；如果直接傳 "FFFFFF" 也可以
        rPr.append(color_el)

    # 隱藏文字 (預設 Word 不顯示)
    if hidden:
        vanish = OxmlElement('w:vanish')
        rPr.append(vanish)

    r.append(rPr)

    # 內容本身只是 placeholder，Word 會自己更新
    t = OxmlElement('w:t')
    t.text = "1"
    r.append(t)

    fld.append(r)
    paragraph._p.append(fld)
    return paragraph

def add_total_minus_one_field(paragraph, font_size=10.5, color=None):
    """
    在 paragraph 中插入 Word 數學公式： { = { NUMPAGES } - 1 }
    font_size: 字體大小
    color: RGBColor(...) 或 16 進位字串 (例如 "FFFFFF" 代表白色)
    """
    p = paragraph._p
    
    # 輔助函式：建立一個 run 並設定字體大小與顏色
    def create_r():
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # 設定字型大小
        if font_size is not None:
            val = str(int(font_size * 2))
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), val)
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), val)
            rPr.append(szCs)
            
        # 設定字型顏色
        if color is not None:
            color_el = OxmlElement('w:color')
            # 支援傳入 RGBColor 物件或是直接傳字串 "FFFFFF"
            color_val = f"{color.rgb:06X}" if hasattr(color, "rgb") else str(color)
            color_el.set(qn('w:val'), color_val)
            rPr.append(color_el)
            
        # 只有在有設定樣式時，才把 rPr 加入 run 裡面
        if len(rPr) > 0:
            r.append(rPr)
            
        return r

    # 1. 外部公式開始 { = 
    r_begin = create_r()
    r_begin.append(OxmlElement('w:fldChar'))
    r_begin.find(qn('w:fldChar')).set(qn('w:fldCharType'), 'begin')
    p.append(r_begin)

    r_instr_eq = create_r()
    instr_eq = OxmlElement('w:instrText')
    instr_eq.set(qn('xml:space'), 'preserve')
    instr_eq.text = ' = '
    r_instr_eq.append(instr_eq)
    p.append(r_instr_eq)

    # 2. 內部 NUMPAGES 開始 { NUMPAGES }
    r_num_begin = create_r()
    r_num_begin.append(OxmlElement('w:fldChar'))
    r_num_begin.find(qn('w:fldChar')).set(qn('w:fldCharType'), 'begin')
    p.append(r_num_begin)

    r_num_instr = create_r()
    num_instr = OxmlElement('w:instrText')
    num_instr.text = 'NUMPAGES'
    r_num_instr.append(num_instr)
    p.append(r_num_instr)

    r_num_end = create_r()
    r_num_end.append(OxmlElement('w:fldChar'))
    r_num_end.find(qn('w:fldChar')).set(qn('w:fldCharType'), 'end')
    p.append(r_num_end)

    # 3. 減法運算 - 1
    r_minus = create_r()
    minus_instr = OxmlElement('w:instrText')
    minus_instr.set(qn('xml:space'), 'preserve')
    minus_instr.text = ' - 1 '
    r_minus.append(minus_instr)
    p.append(r_minus)

    # 4. 公式分隔符號與結束
    r_sep = create_r()
    r_sep.append(OxmlElement('w:fldChar'))
    r_sep.find(qn('w:fldChar')).set(qn('w:fldCharType'), 'separate')
    p.append(r_sep)

    # 預設顯示值 (Word 打開後會自動更新)
    r_val = create_r()
    t = OxmlElement('w:t')
    t.text = '1'
    r_val.append(t)
    p.append(r_val)

    r_end = create_r()
    r_end.append(OxmlElement('w:fldChar'))
    r_end.find(qn('w:fldChar')).set(qn('w:fldCharType'), 'end')
    p.append(r_end)

def set_section_start_page(section, start_num: int = 0):
    """
    設定這個 section 的起始頁碼。start_num = 0 表示第一頁顯示 0，
    但如果第一頁 header 不顯示頁碼（Different first page），
    從第二頁開始看到的是 1、2、3...
    """
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:start'), str(start_num))

def setup_page_numbers(doc):
    """
    第一頁 header：放一個「隱藏的 NUMPAGES」+ 一個正常顯示的 NUMPAGES
    其他頁 header：放 PAGE
    """
    # 第一頁 header
    first_header = doc.sections[0].header
    if first_header and first_header.tables:
        tbl_first = first_header.tables[0]
        cell_total = tbl_first.rows[1].cells[10]   # 你原本用的那格

        if cell_total.paragraphs:
            p_total = cell_total.paragraphs[0]
        else:
            p_total = cell_total.add_paragraph()

        clear_paragraph(p_total)

        # (1) 隱藏的 NUMPAGES：促使 Word 重算整份文件
        add_total_minus_one_field(p_total, font_size = 12, color = "FFFFFF")
        add_total_minus_one_field(p_total, font_size = 12)
        # add_simple_field(p_total, "NUMPAGES \\* MERGEFORMAT", font_size = 12, color = "FFFFFF")
        # add_simple_field(p_total, "NUMPAGES \\* MERGEFORMAT", font_size = 12)

    # 其他頁 header：顯示頁碼 PAGE
    normal_header = doc.sections[1].header
    if normal_header and normal_header.tables:
        tbl_other = normal_header.tables[0]
        cell_page = tbl_other.rows[1].cells[5]

        if cell_page.paragraphs:
            p_page = cell_page.paragraphs[0]
        else:
            p_page = cell_page.add_paragraph()

        clear_paragraph(p_page)

        add_simple_field(p_page, "PAGE", font_size=12)
        run = p_page.add_run(" / ")
        run.font.size = Pt(10.5)
        add_simple_field(p_page, "SECTIONPAGES \\* MERGEFORMAT", font_size=12)

# Extract plain text and "first seen" color from a cell JSON (your ProseMirror-ish structure).
def _extract_text_and_color_from_cell_json(cell_json, COLOR_DICT):
    txt_parts = []
    first_color = None
    for blk in cell_json.get("content", []):
        if blk.get("type") == "paragraph":
            for t in blk.get("content", []) or []:
                if t.get("type") == "text":
                    txt_parts.append(t.get("text", ""))
                    marks = t.get("marks")
                    if first_color is None and marks:
                        ckey = marks[0].get("attrs", {}).get("color")
                        if ckey in COLOR_DICT:
                            first_color = COLOR_DICT[ckey]
    return "".join(txt_parts).strip(), first_color

# Extract header texts for row0 (works if headers are in tableHeader/tableCell/customTableCell)
def _header_texts(row0_cells):
    headers = []
    for cell in row0_cells:
        if cell.get("type") not in ["customTableCell", "tableCell", "tableHeader"]:
            continue
        htxt = []
        for blk in cell.get("content", []):
            if blk.get("type") == "paragraph":
                for t in blk.get("content", []) or []:
                    if t.get("type") == "text":
                        htxt.append(t.get("text", ""))
        headers.append("".join(htxt).strip())
    return headers

# Find indices of the 5 special headers among ALL headers. Returns dict if all present, else None.
def _find_spec_header_indices(headers):
    pos = {}
    for name in SPEC_HDRS:
        try:
            pos[name] = headers.index(name)
        except ValueError:
            return None
    return pos

# Very tolerant float extractor (find first number in text like "12.3mm" or "±2.0")
_num_pat = re.compile(r"[-+]?\d+(?:\.\d+)?")
def _to_float_or_none(s):
    if s is None: return None
    m = _num_pat.search(str(s))
    return float(m.group(0)) if m else None

_calc_significant_digits = lambda s: 0 if "." not in s else len(s.split(".")[-1])
def _compose_value(set_txt, upper_txt, lower_txt):
    s = _to_float_or_none(set_txt)
    u = _to_float_or_none(upper_txt)
    l = _to_float_or_none(lower_txt)

    if u == None and l == None:
        return ""
    
    elif u == None:
        return f"≥{l}"

    elif l == None:
        return f"≤{u}"

    elif s == None:
        return f"{lower_txt}~{upper_txt}"
    
    elif set_txt == upper_txt and set_txt == lower_txt:
        return set_txt

    u_sd = _calc_significant_digits(upper_txt)
    l_sd = _calc_significant_digits(lower_txt)
    s_sd = _calc_significant_digits(set_txt)
    plus = round(u - s, max(s_sd, u_sd)) if s_sd + u_sd != 0 else int(u - s)
    minus = round(s - l, max(s_sd, l_sd)) if s_sd + l_sd != 0 else int(s - l)
    return f"-{minus}/+{plus}\n({lower_txt}~{upper_txt})"

def set_style_fonts(style, latin="Arial", east_asia="標楷體"):
    """
    Apply Latin (ascii/hAnsi/cs) and East Asian (eastAsia) fonts to a Word style.
    Works for paragraph and character styles.
    """
    # Set high-level python-docx font (affects ascii/hAnsi sometimes, but we set XML explicitly too)
    if hasattr(style, "font"):
        style.font.name = latin

    # Ensure <w:rPr> exists
    rPr = style._element.rPr
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        style._element.append(rPr)

    # Ensure <w:rFonts> exists
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)

    # Set Latin/complex + East Asian
    rFonts.set(qn("w:ascii"), latin)          # Latin
    rFonts.set(qn("w:hAnsi"), east_asia)          # Latin (alternate)
    rFonts.set(qn("w:cs"), latin)             # Complex script (safe to mirror Latin)
    rFonts.set(qn("w:eastAsia"), east_asia)   # East Asian (Chinese)

def apply_default_fonts(doc, latin="Arial", east_asia="標楷體"):
    # Normal
    set_style_fonts(doc.styles["Normal"], latin, east_asia)

    # Headings 1..9 (apply if present)
    for i in range(1, 10):
        name = f"Heading {i}"
        if name in doc.styles:
            set_style_fonts(doc.styles[name], latin, east_asia)

    # Common extra styles (optional; add any you use)
    for s in ["Title", "Subtitle", "Intense Quote", "Quote", "Strong", "Emphasis"]:
        if s in doc.styles:
            set_style_fonts(doc.styles[s], latin, east_asia)

def style_json_runs(style_json):
    """攤平 tiptap style_json（{type:doc, content:[paragraph...]}）→ [(text, RGBColor)]。
    只取顏色 mark（沿用 COLOR_DICT）；跨段落串接成單行（header / 目的用）。非 dict / 無文字 → []。"""
    out = []
    if not isinstance(style_json, dict):
        return out
    for block in style_json.get("content") or []:
        for item in (block.get("content") or []):
            if item.get("type") == "text" and item.get("text"):
                color = COLOR_DICT["#000"]
                for m in (item.get("marks") or []):
                    ck = (m.get("attrs") or {}).get("color")
                    if ck in COLOR_DICT:
                        color = COLOR_DICT[ck]
                out.append((item["text"], color))
    return out

def fill_paragraph_styled(p, style_json):
    """把佔位段落（如 header 的 DOC_NAME/PROJECT）換成 style_json 的彩色 runs（只加顏色、沿用原字級）。
    有實際文字 → 渲染並回 True；無 → 不動回 False（呼叫端 fallback 純文字）。"""
    parts = style_json_runs(style_json)
    if not parts:
        return False
    size = p.runs[0].font.size if (p.runs and p.runs[0].font.size) else None
    for r in p.runs:
        r.text = ""
    for text, color in parts:
        for run in set_table_run_text(p, text, color):
            if size:
                run.font.size = size
    return True

def update_p(p, mapping, styled_mapping=None):
    if len(p.runs) == 0:
        return

    full_text = "".join([pr.text for pr in p.runs])

    if mapping.get(full_text) == None:
        return

    if ("POINT" in full_text) or ("REASON" in full_text) or ("DOC_NAME" in full_text) or ("PROJECT" in full_text) or ("ITEM_TYPE" in full_text) or ("STYLE_NO" in full_text):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 彩色覆寫：DOC_NAME / PROJECT 帶 form_attribute style_json 時，注入彩色 runs（無文字則 fallback 純文字）
    if styled_mapping and styled_mapping.get(full_text) and fill_paragraph_styled(p, styled_mapping[full_text]):
        return

    for r in p.runs:
        r.text = ""

    p.runs[0].text = str(mapping[full_text])
    set_run_node_text_font_style(p.runs[0])

def replace_in_footer(footer, title_mapping):
    for row in footer.tables[0].rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                update_p(p, title_mapping)

def replace_in_header(header, title_mapping, styled_mapping=None):
    for row in header.tables[0].rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                is_title_paragraph = "TITLE" in p.text
                update_p(p, title_mapping, styled_mapping)
                if is_title_paragraph:
                    p.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE

def set_repeat_table_header(row):
    """ set repeat table row on every new page"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row

# --- New Content Generation Helper Functions ---
def set_run_node_text_font_style(run_node, color = COLOR_DICT['black']):
    # 基本字型：英文 Arial，中文 標楷體
    run_node.font.name = "Arial"

    r = run_node._element  # CT_R
    # ⭐ 正確取得/建立 <w:rPr>
    rPr = r.get_or_add_rPr()

    # ⭐ 取得/建立 <w:rFonts>
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)

    # 設定字型：英文 + 中文
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rFonts.set(qn("w:eastAsia"), "標楷體")

    if (COLOR_DICT['black'] != color):
        run_node.font.color.rgb = color
        run_node.underline = True

def set_table_run_text(p_node, text, color = COLOR_DICT['#000']):
    runs = []
    if not text:
        return runs
        
    # 利用 Regex 將特殊符號切分出來 (保留符號本身)
    parts = re.split(r'([■□℃≦≧≒℉])', text)
    for part in parts:
        if not part:
            continue
            
        run = p_node.add_run(part)
        run.font.color.rgb = color
        
        if part in ["■", "□", "℃", "℉", "≦", "≧", "≒"]:
            # 針對這兩個符號，強制全面設定為標楷體 (突破 Word 的強制英文渲染)
            run.font.name = "標楷體"
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:ascii"), "標楷體")
            rFonts.set(qn("w:hAnsi"), "標楷體")
            rFonts.set(qn("w:eastAsia"), "標楷體")
            rFonts.set(qn("w:cs"), "標楷體")

            rFonts.set(qn("w:hint"), "eastAsia")
        else:
            # 一般文字走原本的邏輯 (英文 Arial, 中文 標楷體)
            set_run_node_text_font_style(run)
            
        # 如果是藍色，自動加底線
        if color == COLOR_DICT["#0000ff"]:
            run.underline = True
            
        runs.append(run)
        
    # 回傳所有的 run 陣列，方便外部統一處理(如字體大小)
    return runs

def set_docx_table_cell_text(cell_node, text, color = COLOR_DICT["#000"], vCenter = False, center = False):
    p = cell_node.paragraphs[0]
    set_table_run_text(p, text, color)

    if vCenter:
        cell_node.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    return p

# ----- helper function for table merge processing and cell filling ----- #
def _extract_tiptap_table(json_table_data):
    """
    接受整個 doc 或 table node,
    回傳 table_node 與 rows_data(list of tableRow)
    """
    node_type = json_table_data.get("type")
    if node_type == "doc":
        # doc -> 第一個 table
        table_node = next(
            (n for n in json_table_data.get("content", []) if n.get("type") == "table"),
            None
        )
    elif node_type == "table":
        table_node = json_table_data
    else:
        table_node = None

    if not table_node:
        return None, []

    rows_data = [r for r in table_node.get("content", []) if r.get("type") == "tableRow"]
    return table_node, rows_data

def _compute_table_size(rows_data):
    """
    根據 TipTap 的 colspan / rowspan 計算邏輯上的 rowCount, colCount
    colCount = 每一 row 的 colspan 加總的最大值
    rowCount = rows_data 的長度
    """
    rowCount = len(rows_data)
    colCount = 0

    for row in rows_data:
        cells = [c for c in row.get("content", []) if c.get("type") in ["customTableCell", "tableCell", "tableHeader"]]
        col_span_sum = 0
        for cell in cells:
            attr = cell.get("attrs", {}) or {}
            col_span_sum += int(attr.get("colspan", 1) or 1)
        colCount = max(colCount, col_span_sum)

    return rowCount, colCount

def _build_span_grid(rows_data, rowCount, colCount):
    """
    建一個 rowCount x colCount 的矩陣 grid
    grid[r][c] = {
       "is_anchor": bool,
       "cell_data": cell_json or None,
       "rowspan": int,
       "colspan": int,
    }
    非 anchor 的被覆蓋格子也會填一格 {is_anchor: False, ...}
    """
    # 初始化為 None
    grid = [[None for _ in range(colCount)] for _ in range(rowCount)]

    for r_idx, row in enumerate(rows_data):
        cells = [c for c in row.get("content", []) if c.get("type") in ["customTableCell", "tableCell", "tableHeader"]]

        c_idx = 0  # 目前要填入的 column index
        for cell in cells:
            # 找到下一個空的 column
            while c_idx < colCount and grid[r_idx][c_idx] is not None:
                c_idx += 1
            if c_idx >= colCount:
                break  # 超出，理論上不會發生，防禦一下

            attr = cell.get("attrs", {}) or {}
            colspan = int(attr.get("colspan", 1) or 1)
            rowspan = int(attr.get("rowspan", 1) or 1)

            # 錨點
            grid[r_idx][c_idx] = {
                "is_anchor": True,
                "cell_data": cell,
                "rowspan": rowspan,
                "colspan": colspan,
            }

            # 填入被覆蓋的格子（非錨點）
            for dr in range(rowspan):
                for dc in range(colspan):
                    rr = r_idx + dr
                    cc = c_idx + dc
                    if rr == r_idx and cc == c_idx:
                        continue  # 錨點已填
                    if rr < rowCount and cc < colCount:
                        grid[rr][cc] = {
                            "is_anchor": False,
                            "cell_data": None,
                            "rowspan": 0,
                            "colspan": 0,
                        }

            c_idx += colspan

    return grid

def _fill_docx_cell_from_tiptap(docx_cell, cell_data, COLOR_DICT, center = True):
    """
    把 TipTap 的 cell_data 內容寫進 python-docx 的 cell
    保留你原本的 dropdown / paragraph / image 邏輯
    """
    attr = cell_data.get("attrs", {}) or {}
    dropdownValue = attr.get("dropdownValue")

    # 清掉 cell 內原本的空段落
    docx_cell.text = ""
    docx_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if dropdownValue:
        p = set_docx_table_cell_text(docx_cell, dropdownValue, color=COLOR_DICT.get(attr.get("dropdownColor")), vCenter=True, center=center)
        p.paragraph_format.keep_together = True
        return

    items = cell_data.get("content", []) or []
    first_para = True

    for item in items:
        if item.get("type") != "paragraph":
            continue

        # 第一個 paragraph 用現有的，之後用 add_paragraph
        p = docx_cell.paragraphs[0] if first_para and docx_cell.paragraphs else docx_cell.add_paragraph()
        first_para = False
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for text_item in item.get("content", []) or []:
            if text_item.get("type") == "text":
                marks = text_item.get("marks")
                runs = set_table_run_text(p, text_item.get("text", ""))
                # run = p.add_run(text_item.get("text", ""))
                # set_run_node_text_font_style(run)
                if marks:
                    # 假設只有一種 color mark
                    color_key = marks[0].get("attrs", {}).get("color")
                    if color_key in COLOR_DICT:
                        for run in runs:
                            run.font.color.rgb = COLOR_DICT[color_key]
                            if COLOR_DICT[color_key] == RGBColor(0, 0, 255):
                                run.underline = True

            elif text_item.get("type") == "image":
                src = text_item["attrs"]["src"]
                run = p.add_run()
                
                # --- 共用的圖片寬度計算邏輯 ---
                cell_w = getattr(docx_cell, "width", None)
                cell_max_width_cm = (cell_w.cm - 0.5) if (cell_w and cell_w.cm) else 4.0
                cell_max_width_cm = max(cell_max_width_cm, 0.5) # 防呆，避免變成負數

                if src.startswith('http'):
                    local_src = src.split("/", 3)[-1]
                    
                    # 嘗試打開本地圖片，取得真實尺寸
                    try:
                        with Image.open(local_src) as img:
                            dpi = img.info.get('dpi', (96, 96))
                            natural_width_cm = (img.width / dpi[0]) * 2.54
                        # 核心邏輯：取「原圖寬度」與「儲存格最大容許寬度」中較小的值
                        final_width_cm = min(natural_width_cm, cell_max_width_cm)
                    except Exception:
                        # 若讀取失敗，就安全地降級使用最大容許寬度
                        final_width_cm = cell_max_width_cm
                        
                    if os.path.exists(local_src):
                        run.add_picture(local_src, width=Cm(final_width_cm))   # 防呆：圖檔不存在則略過
                    
                elif src.startswith('data:image'):
                    base64_data = src.split(",", 1)[-1]
                    image_bytes = base64.b64decode(base64_data)
                    image_stream = io.BytesIO(image_bytes)
                    
                    with Image.open(image_stream) as img:
                        # 網頁截圖通常預設為 96 DPI
                        dpi = img.info.get('dpi', (96, 96))
                        # 算出圖片原本「真實」的物理寬度 (公分)
                        natural_width_cm = (img.width / dpi[0]) * 2.54
                    
                    # 讀完尺寸後，必須把記憶體檔案指標歸零
                    image_stream.seek(0)
                    
                    # 核心邏輯：取「原圖寬度」與「儲存格最大容許寬度」中較小的值
                    final_width_cm = min(natural_width_cm, cell_max_width_cm)
                    
                    # 寫入 Word
                    run.add_picture(image_stream, width=Cm(final_width_cm))


def create_docx_table(cell, json_table_data):
    rows_data = [row for row in json_table_data.get("content", []) if row.get("type") == "tableRow"]
    if not rows_data:
        return

    # Build list of header cells JSON (row 0)
    row0_cells_json = [c for c in rows_data[0].get("content", []) if c.get("type") in ["customTableCell", "tableCell", "tableHeader"]]
    if not row0_cells_json:
        return

    # Get header texts (full list, including extra headers like 項次/槽體/說明 etc.)
    headers = _header_texts(row0_cells_json)
    # Find positions of the 5 special headers; if not all present -> fall back to original rendering
    pos = _find_spec_header_indices(headers)

    # If NOT a spec table, render as-is (your original behavior)
    if pos is None:
        # 1) 抓 table / rows
        table_node, rows_data = _extract_tiptap_table(json_table_data)
        if not rows_data:
            return

        # 2) 計算 rowCount / colCount
        rowCount, colCount = _compute_table_size(rows_data)

        # 3) 建立 span grid
        grid = _build_span_grid(rows_data, rowCount, colCount)

        # 4) 建立 docx table（一次給足 rows / cols）
        table = cell.add_table(rows=rowCount, cols=colCount)
        table.style = "Table Grid"
        # 如果你有想控制寬度，可以在外面處理，這裡先略過
        # table.width = cell.width.cm - Cm(1)

        # 設定 header repeat（第一 row 視為 header）
        if rowCount > 0:
            set_repeat_table_header(table.rows[0])

        # 5) 先填內容，同時記錄 spanMap
        spanMap = {}  # (r,c) -> (rowspan, colspan)

        for r in range(rowCount):
            for c in range(colCount):
                info = grid[r][c]
                if not info or not info.get("is_anchor"):
                    continue  # 非錨點 cell，不用填內容，只等之後 merge

                cell_data = info["cell_data"]
                rowspan = int(info.get("rowspan", 1) or 1)
                colspan = int(info.get("colspan", 1) or 1)

                docx_cell = table.rows[r].cells[c]

                # 寫入實際內容
                _fill_docx_cell_from_tiptap(docx_cell, cell_data, COLOR_DICT)

                # 有合併才記錄
                if rowspan > 1 or colspan > 1:
                    spanMap[(r, c)] = (rowspan, colspan)

        # 6) 全部內容填完，再進行 merge
        for (r, c), (rowspan, colspan) in spanMap.items():
            first = table.cell(r, c)
            last = table.cell(r + rowspan - 1, c + colspan - 1)
            first.merge(last)

        # 🚀 7) 禁止表格列 (Row) 跨頁斷行
        for row in table.rows:
            trPr = row._tr.get_or_add_trPr()
            if trPr.find(qn('w:cantSplit')) is None:
                cantSplit = OxmlElement('w:cantSplit')
                cantSplit.set(qn('w:val'), 'true')
                trPr.append(cantSplit)

        return

    # ---------- SPEC TABLE TRANSFORMATION ----------
    # Determine insertion position = the smallest index among the five headers
    insert_at = min(pos.values())
    # The set of indices to remove (the five spec headers)
    spec_idx_set = set(pos[h] for h in SPEC_HDRS)

    # 若 header 含「定值項目」，連同此欄一併略過 (column 數會少 1)
    if "定值項目" in headers:
        spec_idx_set.add(headers.index("定值項目"))

    # Build new header sequence: copy original headers left->right, but when we hit insert_at, insert ["規格值","操作值"] once and skip all five spec columns.
    new_header_titles = []
    col_map = []  # For non-spec columns: holds original index; for inserted pair: ('SPECPAIR', None)
    for index, header in enumerate(headers):
        # Replace SPEC column to calculated results
        if index == insert_at:
            col_map.append(("SPECPAIR", None, None))
            new_header_titles += ["設定值", "操作值", "規格值"]

        # Skip SPEC column expect first replaced one
        elif index in spec_idx_set:
            continue

        # Keep origin headers, if not SPEC column
        else:
            col_map.append(("COPY", index, len(new_header_titles)))
            new_header_titles.append(header)

    # If insert_at > len(headers) (theoretically impossible), guard
    if not new_header_titles:
        # fallback to original render if something odd happened
        return

    # Create the new docx table with computed column count
    table = cell.add_table(rows=0, cols=len(new_header_titles))
    # table.width = cell.width.cm - Cm(1)
    table.width = cell.sections[-1].page_width - Cm(1)
    table.style = 'Table Grid'

    # Write new header row
    hdr = table.add_row()
    set_repeat_table_header(hdr)
    for j, title in enumerate(new_header_titles):
        hdr.cells[j].text = title
        set_run_node_text_font_style(hdr.cells[j].paragraphs[0].runs[0])
        hdr.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr.cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Prepare a reverse map for quick lookup
    idx_reg_upper = pos["規格上限(OOS+)"]
    idx_reg_lower = pos["規格下限(OOS-)"]
    idx_opr_upper = pos["操作上限(OOC+)"]
    idx_opr_lower = pos["操作下限(OOC-)"]
    idx_set = pos["設定值"]

    NumColIndex = None if "項次" not in headers else headers.index('項次')
    previousNum = None if NumColIndex == None else 1

    SlotColIndex = None if "槽體" not in headers else headers.index('槽體')
    lastSlotName = None
    lastSlotRow = None if SlotColIndex == None else 1
    last_col_merges = []
    # print(f"rows_data: {rows_data}")

    # Now write each data row (skip row0 because it was header)
    for rowIndex, row_data in enumerate(rows_data):
        if rowIndex == 0:
            continue

        row = table.add_row()
        row.allow_row_break_across_pages = False
        row_cells_json = [c for c in row_data.get("content", []) if c.get("type") in ["customTableCell", "tableCell", "tableHeader"]]

        # Pre-extract text & color for all original columns (so we can reference easily)
        texts, colors = [], []
        for orig_idx, cjson in enumerate(row_cells_json):
            t, col = _extract_text_and_color_from_cell_json(cjson, COLOR_DICT)
            texts.append(t)
            colors.append(col)

        SlotName = texts[SlotColIndex]

        # Compose spec/oper values + color flags
        spec_val = _compose_value(texts[idx_set], texts[idx_reg_upper], texts[idx_reg_lower])
        oper_val = _compose_value(texts[idx_set], texts[idx_opr_upper], texts[idx_opr_lower])
        spec_is_red = COLOR_DICT['blue'] if (colors[idx_reg_upper] == COLOR_BLUE) or (colors[idx_reg_lower] == COLOR_BLUE) else COLOR_DICT['#000']
        oper_is_red = COLOR_DICT['blue'] if (colors[idx_opr_upper] == COLOR_BLUE) or (colors[idx_opr_lower] == COLOR_BLUE) else COLOR_DICT['#000']

        # Fill the new row respecting col_map
        for colIndex, (tag, dataIndex, dstIndex) in enumerate(col_map):
            # [NEW 2] 檢查是否為最後一欄，並讀取 rowspan
            # 只有當前欄位是最後一欄時才處理
            if colIndex == len(col_map) - 1:
                # 確保 dataIndex 在範圍內 (避免前端資料結構異常導致 crash)
                if dataIndex < len(row_cells_json):
                    curr_json = row_cells_json[dataIndex]
                    attrs = curr_json.get("attrs", {})
                    # 取得 rowspan，預設為 1
                    r_span = int(attrs.get("rowspan", 1) or 1)
                    
                    if r_span > 1:
                        # 記錄合併資訊: (開始列, 結束列, 目標欄位索引)
                        # rowIndex 對應 table.rows 的 index (因為 table 有 header 且 rows_data 有 header，剛好對齊)
                        start_r = rowIndex
                        end_r = rowIndex + r_span - 1
                        last_col_merges.append((start_r, end_r, dstIndex))
                        
            if tag == "SPECPAIR":
                p = set_docx_table_cell_text(row.cells[colIndex + 0], texts[idx_set], color = colors[idx_set], vCenter = True, center = True)
                p.paragraph_format.keep_together = True
                p = set_docx_table_cell_text(row.cells[colIndex + 1], oper_val, color = oper_is_red, vCenter = True, center = True)
                p.paragraph_format.keep_together = True
                p = set_docx_table_cell_text(row.cells[colIndex + 2], spec_val, color = spec_is_red, vCenter = True, center = True)
                p.paragraph_format.keep_together = True

            elif NumColIndex != None and colIndex == NumColIndex:
                if SlotName != lastSlotName:
                    p = set_docx_table_cell_text(row.cells[dstIndex], str(previousNum), color = colors[dataIndex], vCenter = True, center = True)
                    p.paragraph_format.keep_together = True
                    previousNum += 1

                    if lastSlotName != None and lastSlotRow != rowIndex - 1:
                        A = table.cell(lastSlotRow, dstIndex)
                        B = table.cell(rowIndex - 1, dstIndex)
                        A.merge(B)

            elif SlotColIndex != None and colIndex == SlotColIndex:
                if SlotName != lastSlotName:
                    p = set_docx_table_cell_text(row.cells[dstIndex], texts[dataIndex], color = colors[dataIndex], vCenter = True, center = True)
                    p.paragraph_format.keep_together = True

                    if lastSlotName != None and lastSlotRow != rowIndex - 1:
                        A = table.cell(lastSlotRow, dstIndex)
                        B = table.cell(rowIndex - 1, dstIndex)
                        A.merge(B)

                    lastSlotRow = rowIndex

            else:
                # 一般欄位寫入 (包含最後一欄的文字寫入)
                if dataIndex < len(texts):
                    p = set_docx_table_cell_text(row.cells[dstIndex], texts[dataIndex], color = colors[dataIndex], vCenter = True, center = True)
                    p.paragraph_format.keep_together = True

        lastSlotName = SlotName

    if lastSlotName != None and rowIndex != lastSlotRow:
        if NumColIndex != None:
            A = table.cell(lastSlotRow, NumColIndex)
            B = table.cell(rowIndex, NumColIndex)
            A.merge(B)
        
        A = table.cell(lastSlotRow, SlotColIndex)
        B = table.cell(rowIndex, SlotColIndex)
        A.merge(B)

    # [NEW 3] 執行最後一欄的合併
    # 這一步放在最後，避免干擾前面的文字寫入
    for start_r, end_r, col_idx in last_col_merges:
        try:
            # 防呆：確保結束列沒有超出表格範圍
            if end_r < len(table.rows):
                cell_start = table.cell(start_r, col_idx)
                cell_end = table.cell(end_r, col_idx)
                cell_start.merge(cell_end)
        except Exception as e:
            print(f"Error merging last column cells ({start_r}-{end_r}): {e}")

    return

def prevent_table_break(table):
    """
    設定表格盡量不跨頁 (Keep with next)。
    若表格長度超過一頁，Word 仍會強制換頁 (符合需求)。
    """
    # 1. 設定每一列 (Row) 內部的內容不要斷開 (Allow row to break across pages = False)
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        # 設定 <w:cantSplit>
        cantSplit = trPr.find(qn('w:cantSplit'))
        if cantSplit is None:
            cantSplit = OxmlElement('w:cantSplit')
            # 將 cantSplit 設為 on，表示「禁止」列內容跨頁
            cantSplit.set(qn('w:val'), 'on') 
            trPr.append(cantSplit)

    # 2. 設定「列與列之間」黏在一起 (Keep with next)
    # 邏輯：除了「最後一列」之外，每一列裡面的每一個段落都要設定 keep_with_next = True
    # 這樣 Row 1 會黏 Row 2, Row 2 黏 Row 3... 直到表格結束
    for row in table.rows[:-1]: # 排除最後一列
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True

def keep_table_on_one_page(table, keep_final_row = False):
    # 1. 確保單一儲存格內的文字不會被切兩半
    for row in table.rows:
        row.allow_row_break_across_pages = False
        
    # 2. 將「除了最後一列以外」的所有段落，設定與下段同頁
    # 這樣 Word 就會盡全力把整個表格包在同一頁
    for row in table.rows[:-1]:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.keep_with_next = True
                
    # 3. 確保最後一列「沒有」被設定與下段同頁，讓 Word 可以在表格後方自然換頁
    for cell in table.rows[-1].cells:
        for p in cell.paragraphs:
            p.paragraph_format.keep_with_next = keep_final_row

def create_parameter_table(doc, condition_content, parameter_content, info):
    if (condition_content == None or len(condition_content) == 0) and (parameter_content == None or len(parameter_content) == 0):
        return
    
    # print(f"condition_content: {condition_content}")
    # print(f"prarmeter_content: {parameter_content}")
    
    programCode = '、'.join([p["programCode"] for p in info['programs']])
    machines_name = '、'.join([m['name'] for m in info['machines']])
    info_N_rows = 2 if info["step_type"] == 2 else 3

    table = None
    if condition_content != None and len(condition_content) != 0:
        tableNode = condition_content
        rows_data = [row for row in tableNode['content'][0]['content'] if row.get("type") == "tableRow"]
        cols = len(rows_data[0]['content']) - 1

        table = doc.add_table(rows = info_N_rows + len(rows_data) + 1, cols = cols)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False

        # total_width = cell.width - Cm(0.2)
        total_width = doc.sections[-1].page_width - Cm(2)

        first_col_width = Inches(1)
        if cols > 1:
            other_col_width = (total_width - first_col_width) // (cols - 1)
            first_col_width = int(total_width - (other_col_width * (cols - 1)))
        else:
            other_col_width = int(total_width - first_col_width)   # 防呆：單欄表避免除以 (cols-1)=0

        for row in table.rows:
            row.allow_row_break_across_pages = False

        for c in range(cols):
            table.columns[c].width = first_col_width if c == 0 else other_col_width
            for table_cell in table.columns[c].cells:
                table_cell.width = first_col_width if c == 0 else other_col_width

        set_docx_table_cell_text(table.cell(0, 0), "程式代碼", center = True)
        table.cell(0, 1).merge(table.cell(0, cols - 1))
        set_docx_table_cell_text(table.cell(0, 1), programCode)
        set_docx_table_cell_text(table.cell(1, 0), "機台名稱", center = True)
        table.cell(1, 1).merge(table.cell(1, cols - 1))
        set_docx_table_cell_text(table.cell(1, 1), machines_name)
        set_docx_table_cell_text(table.cell(2, 0), "條件參數", center = True)
        table.cell(2, 0).merge(table.cell(2, cols - 1))

        for rowIndex, row_data in enumerate(rows_data):
            for colIndex, col_data in enumerate(row_data['content'][1:]):
                _fill_docx_cell_from_tiptap(table.cell(rowIndex + 3, colIndex), col_data, COLOR_DICT)

        # prevent_table_break(table)
        keep_table_on_one_page(table, True)

    write_info = (table == None)
    if parameter_content != None and len(parameter_content) != 0:
        tableNode = parameter_content
        rows_data = [row for row in tableNode['content'][0]['content'] if row.get("type") == "tableRow"]
        # 若 header 含「定值項目」該欄將被略過 → cols 需多 -1
        _hdr_preview = _header_texts(rows_data[0]['content'])
        _has_fixed_value = "定值項目" in _hdr_preview
        cols = len(rows_data[0]['content']) - 2 - (1 if _has_fixed_value else 0)

        if condition_content != None and len(condition_content) != 0:
            p_to_remove = doc.paragraphs[-1]._element
            p_to_remove.getparent().remove(p_to_remove)
        
        table = doc.add_table(rows = info_N_rows + len(rows_data) + 1 if write_info else len(rows_data) + 1, cols = cols)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False

        # total_width = cell.width - Cm(0.2)
        total_width = doc.sections[-1].page_width - Cm(2)

        first_col_width = Inches(1)
        if cols > 1:
            other_col_width = (total_width - first_col_width) // (cols - 1)
            first_col_width = int(total_width - (other_col_width * (cols - 1)))
        else:
            other_col_width = int(total_width - first_col_width)   # 防呆：單欄表避免除以 (cols-1)=0

        for row in table.rows:
            row.allow_row_break_across_pages = False

        for c in range(len(table.columns)):
            table.columns[c].width = first_col_width if c == 0 else other_col_width
            for table_cell in table.columns[c].cells:
                table_cell.width = first_col_width if c == 0 else other_col_width

        if write_info:
            set_docx_table_cell_text(table.cell(0, 0), "程式代碼", center = True)
            table.cell(0, 1).merge(table.cell(0, cols - 1))
            set_docx_table_cell_text(table.cell(0, 1), programCode)
            set_docx_table_cell_text(table.cell(1, 0), "機台名稱", center = True)
            table.cell(1, 1).merge(table.cell(1, cols - 1))
            set_docx_table_cell_text(table.cell(1, 1), machines_name)
            if info["step_type"] == 5:
                set_docx_table_cell_text(table.cell(2, 0), "流程順序", center = True)
                table.cell(2, 1).merge(table.cell(2, cols - 1))
                set_docx_table_cell_text(table.cell(2, 1), "、".join([f"第{order}次" for order in info["processOrder"]]))

        if info["step_type"] == 2:
            set_docx_table_cell_text(table.cell(2 if write_info else 0, 0), "製造參數", center = True)
            table.cell(2 if write_info else 0, 0).merge(table.cell(2 if write_info else 0, cols - 1))
        else:
            set_docx_table_cell_text(table.cell(3 if write_info else 0, 0), "製造參數", center = True)
            table.cell(3 if write_info else 0, 0).merge(table.cell(3 if write_info else 0, cols - 1))

        headers = _header_texts(rows_data[0]['content'])
        pos = _find_spec_header_indices(headers)

         # ---------- SPEC TABLE TRANSFORMATION ----------
        # Determine insertion position = the smallest index among the five headers
        insert_at = min(pos.values())
        # The set of indices to remove (the five spec headers)
        spec_idx_set = set(pos[h] for h in SPEC_HDRS)

        # 若 header 含「定值項目」，連同此欄一併略過 (column 數會少 1)
        if "定值項目" in headers:
            spec_idx_set.add(headers.index("定值項目"))

        # Build new header sequence: copy original headers left->right, but when we hit insert_at, insert ["規格值","操作值"] once and skip all five spec columns.
        new_header_titles = []
        col_map = []  # For non-spec columns: holds original index; for inserted pair: ('SPECPAIR', None)
        for index, header in enumerate(headers):
            # Replace SPEC column to calculated results
            if index == insert_at:
                col_map.append(("SPECPAIR", None, None))
                new_header_titles += ["設定值", "操作值", "規格值"]

            # Skip SPEC column expect first replaced one
            elif index in spec_idx_set:
                continue

            # Keep origin headers, if not SPEC column
            else:
                col_map.append(("COPY", index, len(new_header_titles)))
                new_header_titles.append(header)
        
        # Write new header row
        for j, title in enumerate(new_header_titles):
            table.cell(info_N_rows + 1 if write_info else 1, j).text = title
            set_run_node_text_font_style(table.cell(info_N_rows + 1 if write_info else 1, j).paragraphs[0].runs[0])
            table.cell(info_N_rows + 1 if write_info else 1, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(info_N_rows + 1 if write_info else 1, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Prepare a reverse map for quick lookup
        idx_reg_upper = pos["規格上限(OOS+)"]
        idx_reg_lower = pos["規格下限(OOS-)"]
        idx_opr_upper = pos["操作上限(OOC+)"]
        idx_opr_lower = pos["操作下限(OOC-)"]
        idx_set = pos["設定值"]

        NumColIndex = None if "項次" not in headers else headers.index('項次')
        previousNum = None if NumColIndex == None else 1

        ExplainColIndex = None if "說明" not in headers else headers.index('說明')

        SlotColIndex = None if "槽體" not in headers else headers.index('槽體')
        lastSlotName = None
        lastSlotRow = None if SlotColIndex == None else 2
        lastSlotRow = lastSlotRow + info_N_rows if write_info else lastSlotRow

        for rowIndex, row_data in enumerate(rows_data):
            if rowIndex == 0:
                continue

            row = table.rows[rowIndex + info_N_rows + 1 if write_info else rowIndex + 1]
            row.allow_row_break_across_pages = False
            row_cells_json = [c for c in row_data.get("content", []) if c.get("type") in ["customTableCell", "tableCell", "tableHeader"]]

            # Pre-extract text & color for all original columns (so we can reference easily)
            texts, colors = [], []
            for orig_idx, cjson in enumerate(row_cells_json):
                t, col = _extract_text_and_color_from_cell_json(cjson, COLOR_DICT)
                texts.append(t)
                colors.append(col)

            SlotName = texts[SlotColIndex]

            # Compose spec/oper values + color flags
            spec_val = _compose_value(texts[idx_set], texts[idx_reg_upper], texts[idx_reg_lower])
            oper_val = _compose_value(texts[idx_set], texts[idx_opr_upper], texts[idx_opr_lower])
            spec_is_red = COLOR_DICT['blue'] if (colors[idx_reg_upper] == COLOR_BLUE) or (colors[idx_reg_lower] == COLOR_BLUE) else COLOR_DICT['#000']
            oper_is_red = COLOR_DICT['blue'] if (colors[idx_opr_upper] == COLOR_BLUE) or (colors[idx_opr_lower] == COLOR_BLUE) else COLOR_DICT['#000']

            rowIndex = rowIndex + info_N_rows + 1 if write_info else rowIndex + 1

            # Fill the new row respecting col_map
            for colIndex, (tag, dataIndex, dstIndex) in enumerate(col_map):
                if tag == "SPECPAIR":
                    p = set_docx_table_cell_text(row.cells[colIndex + 0], texts[idx_set], color = colors[idx_set], vCenter = True, center = True)
                    p.paragraph_format.keep_together = True
                    p = set_docx_table_cell_text(row.cells[colIndex + 1], oper_val, color = oper_is_red, vCenter = True, center = True)
                    p.paragraph_format.keep_together = True
                    p = set_docx_table_cell_text(row.cells[colIndex + 2], spec_val, color = spec_is_red, vCenter = True, center = True)
                    p.paragraph_format.keep_together = True

                elif NumColIndex != None and dataIndex == NumColIndex:
                    if SlotName != lastSlotName:
                        p = set_docx_table_cell_text(row.cells[dstIndex], str(previousNum), color = colors[dataIndex], vCenter = True, center = True)
                        p.paragraph_format.keep_together = True
                        previousNum += 1

                        if lastSlotName != None and lastSlotRow != rowIndex - 1:
                            A = table.cell(lastSlotRow, dstIndex)
                            B = table.cell(rowIndex - 1, dstIndex)
                            A.merge(B)

                elif SlotColIndex != None and dataIndex == SlotColIndex:
                    if SlotName != lastSlotName:
                        p = set_docx_table_cell_text(row.cells[dstIndex], texts[dataIndex], color = colors[dataIndex], vCenter = True, center = True)
                        p.paragraph_format.keep_together = True

                        if lastSlotName != None and lastSlotRow != rowIndex - 1:
                            A = table.cell(lastSlotRow, dstIndex)
                            B = table.cell(rowIndex - 1, dstIndex)
                            A.merge(B)

                        lastSlotRow = rowIndex

                elif ExplainColIndex != None and dataIndex == ExplainColIndex:
                    p = _fill_docx_cell_from_tiptap(row.cells[dstIndex], row_cells_json[dataIndex], COLOR_DICT = COLOR_DICT, center = False)

                else:
                    p = set_docx_table_cell_text(row.cells[dstIndex], texts[dataIndex], color = colors[dataIndex], vCenter = True, center = True)
                    p.paragraph_format.keep_together = True

            lastSlotName = SlotName

        if lastSlotName != None and rowIndex != lastSlotRow:
            if NumColIndex != None:
                A = table.cell(lastSlotRow, NumColIndex)
                B = table.cell(rowIndex, NumColIndex)
                A.merge(B)
            
            A = table.cell(lastSlotRow, SlotColIndex)
            B = table.cell(rowIndex, SlotColIndex)
            A.merge(B)
            
        # prevent_table_break(table)
        keep_table_on_one_page(table)
    
    elif info["step_type"] == 2:
        set_docx_table_cell_text(table.cell(3 if write_info else 0, 0), "請依照「3. 管理條件」進行製造參數設定與確認。", center = True)
        table.cell(3 if write_info else 0, 0).merge(table.cell(3 if write_info else 0, cols - 1))
    elif info["step_type"] == 5:
        set_docx_table_cell_text(table.cell(4 if write_info else 0, 0), "請依照「2. 製作條件規範」進行製造參數設定與確認。", center = True)
        table.cell(4 if write_info else 0, 0).merge(table.cell(4 if write_info else 0, cols - 1))

def set_cell_border(cell, **kwargs):
    """
    設定儲存格邊框的通用函式
    用法:
    set_cell_border(
        cell, 
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"}, 
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"val": "nil"}, # start 即 left (左邊框)
        end={"val": "nil"},   # end 即 right (右邊框)
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # 檢查 tcBorders 是否存在，不存在則建立
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # 處理傳入的邊框設定
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV', 'start', 'end'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # 如果該邊框設定已存在，先移除舊的
            element = tcBorders.find(qn(tag))
            if element is not None:
                tcBorders.remove(element)

            # 建立新的邊框節點
            border = OxmlElement(tag)
            
            # 設定屬性 (sz=大小, val=樣式, color=顏色, space=間距)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    border.set(qn('w:{}'.format(key)), str(edge_data[key]))
            
            tcBorders.append(border)

def create_process_table(cell, content_data):
    if content_data.get('files') != None and len(content_data.get('files')) > 0:
        # width = cell.width.cm
        width_cm = cell.sections[-1].page_width.cm
        src = content_data['files'][0]['path']
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(f'uploads/{src}', width = Cm(width_cm - 3))
        return
        
    if type(content_data.get('table_json')) != dict or content_data['table_json'].get('content') == None or len(content_data['table_json']['content']) == 0:
        p_attr = cell.add_paragraph("NA")
        p_attr.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_attr.runs[0].font.size = Pt(12)
        p_attr.paragraph_format.left_indent = p_attr.runs[0].font.size * 2
        set_run_node_text_font_style(p_attr.runs[0])
        return 
    
    tableNode = content_data['table_json']['content'][0]
    rows_data = [row for row in tableNode['content'] if row.get("type") == "tableRow"]
    if not rows_data:
        return []
    table = cell.add_table(rows=len(rows_data), cols=len(rows_data[0]['content']))
    table.width = cell.sections[-1].page_width - Cm(1)
    table.style = 'Table Grid'

    for r, row_data in enumerate(rows_data):
        for c, cell_data in enumerate(row_data['content']):
            _fill_docx_cell_from_tiptap(table.rows[r].cells[c], cell_data, COLOR_DICT)

            if cell_data['content'][0].get('content') == None:
                set_cell_border(table.cell(r, c), left = {"val": "nil"}, top = {"val": "nil"}, right = {"val": "nil"}, bottom = {"val": "nil"})

def parse_json_content(parent_object, json_data, indent = None, header = False, no = None, tier = 1, keep_with_next = False):
    """
    Recursively parses the JSON content (header_json/content_json) and adds elements 
    (paragraphs, list items) to the parent object (docx.table._Cell).
    Applies tier-based indentation.
    """
    if json_data is None or json_data.get("content") is None:
        return

    base_indent = Cm(0.7) 

    for block in json_data["content"]:
        block_type = block.get("type")
        
        # Use the 'tier_no' from the content object, or the current recursion tier.
        content_tier = block.get("tier_no", tier) 

        if block_type == "paragraph" or block_type == "listItem":
            p = parent_object.add_paragraph()
            p.style = 'List Paragraph' if block_type == "listItem" else 'Normal'

            if header:
                set_table_run_text(p, f"{no} ")

            # Set text for title in content table
            for item in block.get("content", []):
                if item.get("type") == "text":
                    marks = item.get("marks", [])
                    color = COLOR_DICT["#000"]
                    if len(marks) > 0:
                        color_key = marks[0].get("attrs", {}).get("color")
                        if color_key in COLOR_DICT:
                            color = COLOR_DICT[color_key]
                            # color = COLOR_DICT[marks[0].get("attrs", {"color": "#000"}).get("color")] if len(marks) > 0 else COLOR_DICT["#000"]

                    # run = set_table_run_text(p, item.get("text", ""), color)
                    # set_run_node_text_font_style(run)
                    # run.font.size = Pt(12)

                    runs = set_table_run_text(p, item.get("text", ""), color)
                    for run in runs:
                        run.font.size = Pt(12)
                
                elif item.get("type") == "image":
                    src = item["attrs"]["src"]
                    run = p.add_run()
                    
                    # 定義前端與 Word 的黃金平衡寬度：最大 600 像素
                    MAX_PIXELS = 600

                    if src.startswith('http'):
                        local_src = src.split("/", 3)[-1]
                        
                        try:
                            with Image.open(local_src) as img:
                                dpi = img.info.get('dpi', (96, 96))
                                # 邏輯：取「圖片原像素」與「600 像素」中較小的那個
                                target_px = min(img.width, MAX_PIXELS)
                                # 將像素轉換為 Word 認識的公分 (cm)
                                target_width_cm = (target_px / dpi[0]) * 2.54
                                
                        except Exception as e:
                            print(f"Warning: 無法讀取本地圖片尺寸 {local_src}, error: {e}")
                            # 防呆：萬一讀不到，就給標準的 600px 轉換值 (約 15.875 公分)
                            target_width_cm = (MAX_PIXELS / 96) * 2.54
                            
                        # 寫入 Word
                        if os.path.exists(local_src):
                            run.add_picture(local_src, width=Cm(target_width_cm))   # 防呆：圖檔不存在則略過
                        
                    elif src.startswith('data:image'):
                        # 處理 Base64 圖片
                        base64_data = src.split(",", 1)[-1]
                        image_bytes = base64.b64decode(base64_data)
                        image_stream = io.BytesIO(image_bytes)
                        
                        with Image.open(image_stream) as img:
                            dpi = img.info.get('dpi', (96, 96))
                            target_px = min(img.width, MAX_PIXELS)
                            target_width_cm = (target_px / dpi[0]) * 2.54
                        
                        # 讀完尺寸後，必須把記憶體檔案指標歸零
                        image_stream.seek(0)
                        
                        # 寫入 Word
                        run.add_picture(image_stream, width=Cm(target_width_cm))

                elif item.get("type") == "image":
                    src = item["attrs"]["src"]
                    run = p.add_run()
                    if src.startswith('http'):
                        local_src = src.split("/", 3)[-1]
                        # 以目前 cell 寬度大概估一下圖片寬度
                        cell_width_cm = parent_object.width.cm if hasattr(parent_object, "width") else 4
                        if os.path.exists(local_src):
                            run.add_picture(local_src, width=Cm(max(cell_width_cm - 0.5, 0.5)))   # 防呆：圖檔不存在則略過
                    elif src.startswith('data:image'):
                        # 處理 Base64 圖片 (格式通常是 "data:image/png;base64,iVBORw0K...")
                        # 1. 用逗號切開，取得後面的純 base64 字串
                        base64_data = src.split(",", 1)[-1]
                        
                        # 2. 將 base64 字串解碼成二進位位元組 (bytes)
                        image_bytes = base64.b64decode(base64_data)
                        
                        # 3. 包裝成 BytesIO (記憶體中的檔案物件)
                        image_stream = io.BytesIO(image_bytes)
                        
                        # 4. 直接塞入 Word，不特別設定長寬，讓它維持原始比例
                        run.add_picture(image_stream)


            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            indent_factor = max(0, content_tier - 1)
            p.paragraph_format.left_indent = base_indent * indent_factor

            if keep_with_next:
                p.paragraph_format.keep_with_next = True


def createHeader(cell, content_obj, no=None, depth=2):
    """標題渲染：no 由 DFS 算好傳入；depth 控縮排。header_json 為空但有 no → 仍輸出編號（如 ct=2 PMS 3.1）。"""
    if no is None:
        no = ""
    hj = content_obj.get("header_json")
    if hj:
        parse_json_content(cell, hj, header = True, no = no, tier = depth, keep_with_next = True)
    elif no:
        p = cell.add_paragraph()
        # 指示書 step1 的 3.1 為固定標題「基本生產條件」（metadata.source=management）
        src = (content_obj.get("metadata") or {}).get("source")
        text = f"{no} 基本生產條件" if src == "management" else f"{no} "
        runs = set_table_run_text(p, text)
        for run in runs:
            run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(0.7) * max(0, depth - 1)
        p.paragraph_format.keep_with_next = True

def createContent(cell, content_obj, depth=2):
    """內文渲染：縮排隨 depth（2..8）遞增。"""
    if content_obj.get("content_json"):
        parse_json_content(cell, content_obj["content_json"], tier=depth)

def createPictures(cell, content_obj):
    """Adds picture placeholders/links (files) for all items in the list to the cell."""
    width_cm = cell.sections[-1].page_width.cm
    if content_obj.get("files"):
        for file_info in content_obj["files"]:
            if file_info.get("path_to_save") != None:
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                src = file_info["path_to_save"].split("/", 1)[-1]
                run = p.add_run()
                run.add_picture(f'uploads/temp/{src}', width = Cm(width_cm - 2))

def createTable(cell, content_obj):
    """Adds tables (content_json containing a table) for all items in the list to the cell."""
    if content_obj.get("table_json"):
        # The JSON for table content is expected to be a single 'table' block inside 'content'
        table_blocks = [b for b in content_obj["table_json"].get("content", []) if b.get("type") == "table"]
        for table_block in table_blocks:
            # Table is created directly in the cell, as requested, to avoid left indent.
            create_docx_table(cell, table_block)

def _na_paragraph(cell, text="NA", indent_factor=2):
    """加一行 NA / 提示文字段落。"""
    p = cell.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.runs[0].font.size = Pt(12)
    set_run_node_text_font_style(p.runs[0])
    p.paragraph_format.left_indent = p.runs[0].font.size * indent_factor

def _msg_paragraph(cell, text):
    """加一行固定提示（如 isParamNA 的『請依照…』）。"""
    p = cell.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.runs[0].font.size = Pt(12)
    set_run_node_text_font_style(p.runs[0])
    p.paragraph_format.left_indent = Cm(0.7) * 1

def _render_node_dfs(cell, chapter, index_path, node):
    """
    一般階層節點 DFS 渲染（L2..L8，直接寫進 doc）。順序：標題 → 自身內容 → children（spec §19 F4）。
    """
    depth = 1 + len(index_path)
    no = format_node_number(chapter, index_path)
    option = node.get("content_type")

    if option == 2:        # 標題 + 表格
        createHeader(cell, node, no=no, depth=depth)
        createTable(cell, node)
        createContent(cell, node, depth=depth)
    elif option == 1:      # 標題 + 文字 + 圖
        createHeader(cell, node, no=no, depth=depth)
        createContent(cell, node, depth=depth)
        createPictures(cell, node)
    else:                  # option 0 / 3：純標題
        createHeader(cell, node, no=no, depth=depth)

    for j, child in enumerate(node.get("children") or [], start=1):
        _render_node_dfs(cell, chapter, index_path + [j], child)

def draw_instruction_content(doc, data):
    attribute = data["attribute"][-1]
    content_tree = data["content"]   # build_tree 產出：[{step_type, children}]
    tree_by_step = {int(s["step_type"]): (s.get("children") or []) for s in content_tree}
    references = data["reference"]

    documentType = attribute["document_type"]
    cell = doc

    # Get document chapters from definition
    for (stepIndex, itemInfo) in enumerate(DOCUMENT_STEP[DOCUMENT_TYPE(documentType).name].value):
        (step, stepInfo), = itemInfo.items()
        
        # 1. Set the step title
        if stepIndex > 0:
            cell.add_paragraph()
        p_title = cell.paragraphs[-1] if stepIndex == 0 else cell.add_paragraph()
        p_title.text = f"{stepIndex + 1}.{step}"
        p_title.runs[0].font.size = Pt(12)
        set_run_node_text_font_style(p_title.runs[0])
        p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_title.paragraph_format.keep_with_next = True

        # 2. Get current step content
        if stepInfo["parent"] == "attribute":
            # 目的：優先用 form_attribute 的 style_json（tiptap 彩色）；沒有實際文字 → fallback 純文字
            form_attr = data.get("form_attribute") or {}
            purpose_style = form_attr.get("purpose")   # form_attribute key 固定為 "purpose"（指示書/式樣書皆同）
            if style_json_runs(purpose_style):
                p0 = len(cell.paragraphs)
                parse_json_content(cell, purpose_style)
                for p_attr in cell.paragraphs[p0:]:
                    p_attr.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_attr.paragraph_format.left_indent = Pt(12) * 2
            else:
                stepContent = attribute.get(stepInfo["code"])
                if isinstance(stepContent, str):
                    stepContent = "NA" if len(stepContent) == 0 else stepContent
                    p_attr = cell.add_paragraph(stepContent)
                    p_attr.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_attr.runs[0].font.size = Pt(12)
                    set_run_node_text_font_style(p_attr.runs[0])
                    p_attr.paragraph_format.left_indent = p_attr.runs[0].font.size * 2
        
        elif stepInfo["parent"] == "content":
            # 新階層：依 step_type 取子樹（build_tree 產出），分流渲染
            chapter = stepIndex + 1
            step_type = stepInfo["code"]
            children = tree_by_step.get(step_type, [])
            attr = attribute["attribute"]

            # --- 製造流程圖 (step 0) ---
            if step_type == 0:
                if not children:
                    _na_paragraph(cell)
                else:
                    for i, node in enumerate(children, start=1):
                        createHeader(cell, node, no=format_node_number(chapter, [i]), depth=2)
                        create_process_table(cell, node)
                continue

            # --- 指示書 製造條件參數一覽表 (step 2)：多個 ct=4 兄弟節點 ---
            if step_type == 2:
                if not children or attr.get("isParamNA", False):
                    _msg_paragraph(cell, "請依照「3. 管理條件」進行製造參數設定與確認")
                    continue
                machines_name = attr["inputMachines"].split(",")
                machines_code = attr["machines"]
                machines = [{"code": c, "name": n} for c, n in zip(machines_code, machines_name)]
                for node in children:
                    tj = node.get("table_json") or {}
                    info = {"step_type": 2, "programs": (node.get("metadata") or {}).get("programs", []), "machines": machines, "processOrder": None}
                    create_parameter_table(cell, tj.get("conditionTable"), tj.get("parameterTable"), info)
                continue

            # --- 式樣書 製造參數一覽表 (step 5)：單/多個 ct=4 節點 ---
            if step_type == 5:
                if not children:
                    _na_paragraph(cell)
                    continue
                for node in children:
                    md = node.get("metadata") or {}
                    if md.get("isParamNA", False):
                        _msg_paragraph(cell, "請依照「2. 製作條件規範」進行製造參數設定與確認")
                        continue
                    machines_name = md.get("machines_name", [])
                    machines_code = md.get("machines", [])
                    machines = [{"code": c, "name": n} for c, n in zip(machines_code, machines_name)]
                    tj = node.get("table_json") or {}
                    info = {"step_type": 5, "programs": md.get("programs", []), "machines": machines, "processOrder": md.get("processOrder")}
                    create_parameter_table(cell, None, tj.get("parameterTable"), info)
                continue

            # --- 一般階層 (step 1/3/4/6/7)：DFS L2..L8 ---
            if not children:
                _na_paragraph(cell)
                continue
            for i, node in enumerate(children, start=1):
                _render_node_dfs(cell, chapter, [i], node)

        elif stepInfo["parent"] == "reference":
            # Filter reference items for the current referenceType
            docs = [reference for reference in references if reference["refer_type"] == stepInfo["code"]]
            print(f"docs: {docs}")
            
            if len(docs) > 0:
                # Add a reference list with indentation
                for index, ref in enumerate(docs):
                    # For reference documents, use a simple numbered format
                    p_attr = cell.add_paragraph(f"{stepIndex + 1}.{index + 1} {ref.get('refer_document', '')} - {ref.get('refer_document_name', '')}")
                    p_attr.runs[0].font.size = Pt(12)
                    p_attr.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_attr.paragraph_format.left_indent = p_attr.runs[0].font.size * 2
                    p_attr.paragraph_format.space_after = Pt(3)
                    set_run_node_text_font_style(p_attr.runs[0], COLOR_DICT[ref.get("color", 'black')])

            else:
                docs = "NA" if len(docs) == 0 else docs
                p_attr = cell.add_paragraph(docs)
                p_attr.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_attr.runs[0].font.size = Pt(12)
                p_attr.paragraph_format.left_indent = p_attr.runs[0].font.size * 2
                set_run_node_text_font_style(p_attr.runs[0])

def generate_word_password_hash(password, spin_count=100000):
    """
    依照 Office Open XML (OOXML) 標準產生密碼的 salt 與 hash
    """
    # 1. 產生 16 bytes 的隨機鹽值
    salt = os.urandom(16)
    
    # 2. Word 規定密碼必須使用 UTF-16 LE 編碼
    pwd_bytes = password.encode('utf-16le')
    
    # 3. 初始雜湊：SHA512(salt + password)
    hash_calc = hashlib.sha512(salt + pwd_bytes).digest()
    
    # 4. 進行 Spin 迭代運算 (預設 10 萬次)
    for i in range(spin_count):
        # 迭代器必須是 4 bytes 的 Little-Endian 格式
        iterator = i.to_bytes(4, byteorder='little')
        hash_calc = hashlib.sha512(hash_calc + iterator).digest()
        
    # 5. 回傳 Base64 編碼後的字串
    salt_b64 = base64.b64encode(salt).decode('utf-8')
    hash_b64 = base64.b64encode(hash_calc).decode('utf-8')
    
    return salt_b64, hash_b64

def enable_docx_protection(doc, password):
    """
    直接修改 python-docx 的 document 物件，注入保護設定與密碼驗證參數。
    """
    settings_element = doc.settings.element

    # 取得符合 Word 標準的 salt 與 hash
    salt_value, hash_value = generate_word_password_hash(password)

    protection = OxmlElement('w:documentProtection')
    
    # 設定保護類型
    protection.set(qn('w:edit'), 'forms')  
    protection.set(qn('w:enforcement'), '1')
    
    # 設定加密參數
    protection.set(qn('w:cryptProviderType'), 'rsaAES')
    protection.set(qn('w:cryptAlgorithmClass'), 'hash')
    protection.set(qn('w:cryptAlgorithmType'), 'typeAny')
    protection.set(qn('w:cryptAlgorithmSid'), '14') # 14 代表 SHA-512
    protection.set(qn('w:cryptSpinCount'), '100000')
    
    # ★ 關鍵修正：必須同時寫入 salt 與 hash
    protection.set(qn('w:salt'), salt_value)
    protection.set(qn('w:hash'), hash_value)

    # 移除舊標籤
    existing = settings_element.find(qn('w:documentProtection'))
    if existing is not None:
        settings_element.remove(existing)

    # 寫入新標籤
    settings_element.append(protection)

def fill_from_template(template_path, out_path, data, title_mapping, info_mapping, styled_mapping=None):
    doc = Document(template_path)
    apply_default_fonts(doc, latin="Arial", east_asia="標楷體")

    # 產生內容（影響頁數）
    draw_instruction_content(doc, data)


    # 先做原本的 header/footer 文字替換
    for index, section in enumerate(doc.sections):
        # 如果你的第一頁 header 也要做文字替換，可以另外處理：
        # replace_in_header(section.first_page_header, title_mapping)  # 視需求決定
        replace_in_header(section.header, title_mapping, styled_mapping)
        replace_in_footer(section.footer, title_mapping)
        # print(f"index: {index}, section: {section.header}")

    for row in doc.tables[0].rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                update_p(p, info_mapping)

    # ★ 內容都產生完之後，再插入頁碼欄位
    setup_page_numbers(doc)

    enable_docx_protection(doc, "1q2w3e4R")

    fix_image_id_clash(doc)

    doc.save(out_path)

_code_prefix_re = re.compile(r"^\s*\(([^)]+)\)\s*(.*)$")

def clean_process_name(raw: str) -> str:
    """
    PROCESS_NAME like "(L100-01)單面前處理" -> "單面前處理"
    If there is no "(...)" prefix, returns the original, trimmed.
    """
    if not raw:
        return ""
    m = _code_prefix_re.match(raw)
    return m.group(2).strip() if m else raw

def get_docx_without_framework_(outpath, data, template = "docx-template/example3.docx"):
    attribute = data["attribute"][-1]
    Doc_id = attribute["document_id"]
    # 制定日期：預覽快照時帶入「當初產生文件的日期」(data["render_date"])；產生 Word / 即時預覽則用 now()
    Date = data.get("render_date") or datetime.now().strftime("%Y/%m/%d")
    Version = f"{float(attribute['document_version']):.1f}"
    Doc_name = attribute["document_name"]

    itemType = attribute["attribute"].get("itemType")
    styleNo = attribute["attribute"].get("styleNo")

    # 👇 新增 documentKey（每次下載的唯一 key）
    doc_key = attribute.get("documentKey", "")

    title_mapping = {
        "DOC_NO": Doc_id,
        "DATE": Date,
        "REV": Version,
        "PAGE": "1",
        "DOC_NAME": Doc_name,
        "PROJECT": "",
        "ITEM_TYPE": "",
        "STYLE_NO": "",
        "DOC_KEY": doc_key,
    }
    if itemType is not None:
        specifications = [clean_process_name(s["name"]) for s in attribute["attribute"].get("specification")]
        title_mapping["PROJECT"] = "_".join(specifications)
        title_mapping["ITEM_TYPE"]=  styleNo.split("-", 1)[0]
        title_mapping["STYLE_NO"] = styleNo.split("-", 1)[-1]
    else:
        title_mapping["PROJECT"] = attribute["attribute"].get("applyProject", "")

    # 指示書才用 form_attribute 的彩色 DOC_NAME / PROJECT（式樣書頁首只有純文字；目的另在 body 處理）
    form_attr = data.get("form_attribute") or {}
    styled_mapping = {}
    if attribute.get("document_type", 0) == 0:
        if style_json_runs(form_attr.get("document_name")):
            styled_mapping["DOC_NAME"] = form_attr.get("document_name")
        if style_json_runs(form_attr.get("applyProject")):
            styled_mapping["PROJECT"] = form_attr.get("applyProject")

    info_mapping = {
        "REV1": "", "DATE1": "", "REASON1": "", "POINT1": "", "DEPT1": attribute["department"], "APPROVER1": "", "CONFIRMER1": "", "AUTHOR1": attribute["author"],
        "REV2": "", "DATE2": "", "REASON2": "", "POINT2": "", "DEPT2": "", "APPROVER2": "", "CONFIRMER2": "", "AUTHOR2": "",
        "REV3": "", "DATE3": "", "REASON3": "", "POINT3": "", "DEPT3": "", "APPROVER3": "", "CONFIRMER3": "", "AUTHOR3": ""
    }

    # Place document attribute into Word
    for index, attr in enumerate(data["attribute"]):
        info_mapping[f"REV{index + 1}"] = f'{float(attr["document_version"]):.1f}'
        info_mapping[f"DATE{index + 1}"] = attr.get("issue_date", datetime.now().strftime("%Y/%m/%d"))
        info_mapping[f"DEPT{index + 1}"] = attribute["department"]
        info_mapping[f"APPROVER{index + 1}"] = attr["approver"]
        info_mapping[f"CONFIRMER{index + 1}"] = attr["confirmer"]
        info_mapping[f"AUTHOR{index + 1}"] = attribute["author"]
        info_mapping[f"REASON{index + 1}"] = f'變更理由\n{attr["change_reason"]}'
        info_mapping[f"POINT{index + 1}"] = f'變更要點\n{attr["change_summary"]}'

    # Assuming 'example__.docx' exists in the execution environment
    fill_from_template(template, outpath, data, title_mapping, info_mapping, styled_mapping)

def fix_image_id_clash(document):
    """
    修復 python-docx 插入圖片導致的 wp:docPr ID 衝突問題 (Issue #455)
    這會將主要文件中的所有圖片/圖案 ID 加上 100000，避免與頁首/頁尾的 ID 衝突。
    """
    # 取得整份主要文件的 XML 根節點
    doc_element = document._part._element
    
    # 找出所有圖片的屬性節點 (wp:docPr)
    docPrs = doc_element.findall('.//' + qn('wp:docPr'))
    
    for docPr in docPrs:
        # 取得目前的 ID，並加上 100000
        current_id = int(docPr.get('id', '0'))
        new_id = current_id + 100000
        
        # 重新寫回 XML 中
        docPr.set('id', str(new_id))

# ----------------- Example usage -----------------
if __name__ == "__main__":
    # The file path has been changed to match the payload.json fetched content
    # Note: Replace '_captures/docs/20251028-114010-2b096f8a/payload.json' 
    # with the actual file path if running locally. 
    # Since I have the content, I will use the dictionary directly for demonstration.
    # with open("_captures/docs/20251030-161207-516a0cfa/payload.json", 'r', encoding = "utf-8") as f:
    with open("_captures/docs/20251103-095207-27ee69e9/payload.json", 'r', encoding = "utf-8") as f:
        data = json.load(f)

    # template = r"./example__.docx"   # <- your example file (the “second” screenshot)
    # output   = r"./header_table_fix.docx"
    template = "docx-template/example3.docx"
    output = "docxTemp/temp.docx"

    attribute = data["attribute"][-1]
    Doc_id = attribute["document_id"]
    Date = datetime.now().strftime("%Y/%m/%d")
    Version = f"{int(attribute['document_version']):.1f}"
    Title = "製造條件指示書" if attribute["document_type"] == 0 else "製造式樣書"
    Doc_name = attribute["document_name"]

    # Put placeholders like [DOC_NO], [REV], [DATE], [TOTAL_PAGES] in the header cells of your template.
    # title_mapping = { "DOC_NO": "WMH250", "DATE": "2025/10/29", "REV": "1.0", "PAGE": "1", "TITLE": "製造條件指示書", "DOC_NAME": "", "PROJECT": "", "DOC_CODE": "FM-R-MF-AZ-052 Rev7.0"}
    title_mapping = { "DOC_NO": Doc_id, "DATE": Date, "REV": Version, "PAGE": "1", "TITLE": Title, "DOC_NAME": Doc_name, "PROJECT": "", "DOC_CODE": "FM-R-MF-AZ-052 Rev7.0"}
    info_mapping = {
        "REV1": "", "DATE1": "", "REASON1": "", "POINT1": "", "DEPT1": "", "APPROVER1": "", "CONFIRMER1": "", "AUTHOR1": "",
        "REV2": "", "DATE2": "", "REASON2": "", "POINT2": "", "DEPT2": "", "APPROVER2": "", "CONFIRMER2": "", "AUTHOR2": "",
        "REV3": "", "DATE3": "", "REASON3": "", "POINT3": "", "DEPT3": "", "APPROVER3": "", "CONFIRMER3": "", "AUTHOR3": ""
    }

    for index, attribute in enumerate(data["attribute"]):
        info_mapping[f"REV{index + 1}"] = f'{attribute["document_version"]:.1f}'
        info_mapping[f"DATE{index + 1}"] = attribute.get("issueDate", datetime.now().strftime("%Y/%m/%d"))
        info_mapping[f"REASON{index + 1}"] = attribute["reviseReason"]
        info_mapping[f"POINT{index + 1}"] = attribute["revisePoint"]
        info_mapping[f"DEPT{index + 1}"] = attribute["department"]
        info_mapping[f"APPROVER{index + 1}"] = attribute["approver"]
        info_mapping[f"CONFIRMER{index + 1}"] = attribute["confirmer"]
        info_mapping[f"AUTHOR{index + 1}"] = attribute["author"]
        info_mapping[f"REASON{index + 1}"] = f'變更理由\n{attribute["reviseReason"]}'
        info_mapping[f"POINT{index + 1}"] = f'變更要點\n{attribute["revisePoint"]}'

    # Assuming 'example__.docx' exists in the execution environment
    fill_from_template(template, output, data, title_mapping, info_mapping) 
    print("The script with the completed draw_instruction_content function is ready.")
    print("Please ensure your local environment has 'example__.docx' and python-docx installed to run the full script.")
    