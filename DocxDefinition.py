import json, re, os, base64
import hashlib
from datetime import datetime
from enum import Enum

from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

COLOR_DICT = {
    "red": RGBColor(255, 0, 0),
    "blue": RGBColor(0, 0, 255),
    "#000": RGBColor(0, 0, 0),
    "#000000": RGBColor(0, 0, 0),
    "#ff0000": RGBColor(255, 0, 0),
    "#0000ff": RGBColor(0, 0, 255),
}

class DOCUMENT_TYPE(Enum):
    ManufacturingDocument = 0
    SpecificationDocument = 1

class DOCUMENT_STEP(Enum):
    ManufacturingDocument = [
        {"目的": {"parent": "attribute", "code": "documentPurpose"}},
        {"製造流程": {"parent": "content", "code": 0}},
        {"管理條件": {"parent": "content", "code": 1}},
        {"製造條件參數一覽表": {"parent": "content", "code": 2}},
        {"異常處置": {"parent": "content", "code": 3}},
        {"相關文件": {"parent": "reference", "code": 0}},
        {"使用表單": {"parent": "reference", "code": 1}}
    ]
    SpecificationDocument = [
        {"目的": {"parent": "attribute", "code": "documentPurpose"}},
        {"製作條件規範": {"parent": "content", "code": 4}},
        {"製造參數一覽表": {"parent": "content", "code": 5}},
        {"適用品質與規格內容": {"parent": "content", "code": 6}},
        {"使用表單": {"parent": "reference", "code": 1}},
        {"其他": {"parent": "content", "code": 7}},
    ]


SPEC_HDRS = ['規格下限(OOS-)','操作下限(OOC-)','設定值','操作上限(OOC+)','規格上限(OOS+)']
COLOR_RED = RGBColor(255, 0, 0)
COLOR_BLUE = RGBColor(0, 0, 255)

def clear_paragraph(p):
    """把一個 paragraph 內原本所有 run/文字清掉"""
    p.text = ""
    # 保險一點也可以把底層 element 都清掉
    for r in p.runs:
        r._element.getparent().remove(r._element)

def add_simple_field(paragraph, instr, font_size=None, hidden=False, color=None):
    """
    在 paragraph 中插入一個簡單欄位，例如:
      instr = "PAGE" 或 "NUMPAGES"
    font_size: Pt(...) / 數字 (例如 10.5)，會轉成半點數給 w:sz
    hidden: True => 使用 w:vanish 隱藏這個 run
    color: RGBColor(...)，可選 (例如 RGBColor(255, 255, 255) = 白色)
    """
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), instr)  # e.g. "PAGE", "NUMPAGES"

    r = OxmlElement('w:r')

    # ---- run properties (rPr) ----
    rPr = OxmlElement('w:rPr')

    # 字型大小
    if font_size is not None:
        # Word 的 w:sz 是「半點數」：10.5pt => 21
        val = int(round(float(font_size) * 2))
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(val))
        rPr.append(sz)

        szCs = OxmlElement('w:szCs')  # complex script size，順便設
        szCs.set(qn('w:val'), str(val))
        rPr.append(szCs)

    # 顏色
    if color is not None:
        color_el = OxmlElement('w:color')
        color_el.set(qn('w:val'), f"{color.rgb:06X}" if hasattr(color, "rgb") else color)
        # 如果用 RGBColor，就使用 color.rgb；如果直接傳 "FFFFFF" 也可以
        rPr.append(color_el)

    # 隱藏文字 (預設 Word 不顯示)
    if hidden:
        vanish = OxmlElement('w:vanish')
        rPr.append(vanish)

    r.append(rPr)

    # 內容本身只是 placeholder，Word 會自己更新
    t = OxmlElement('w:t')
    t.text = "1"
    r.append(t)

    fld.append(r)
    paragraph._p.append(fld)
    return paragraph

def set_section_start_page(section, start_num: int = 0):
    """
    設定這個 section 的起始頁碼。start_num = 0 表示第一頁顯示 0，
    但如果第一頁 header 不顯示頁碼（Different first page），
    從第二頁開始看到的是 1、2、3...
    """
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:start'), str(start_num))

def setup_page_numbers(doc):
    """
    第一頁 header：放一個「隱藏的 NUMPAGES」+ 一個正常顯示的 NUMPAGES
    其他頁 header：放 PAGE
    """
    # 第一頁 header
    first_header = doc.sections[0].header
    if first_header and first_header.tables:
        tbl_first = first_header.tables[0]
        cell_total = tbl_first.rows[1].cells[10]   # 你原本用的那格

        if cell_total.paragraphs:
            p_total = cell_total.paragraphs[0]
        else:
            p_total = cell_total.add_paragraph()

        clear_paragraph(p_total)

        # (1) 隱藏的 NUMPAGES：促使 Word 重算整份文件
        add_simple_field(p_total,"NUMPAGES",font_size=1,hidden=True)

        # (2) 正常顯示的 NUMPAGES
        add_simple_field(p_total, "NUMPAGES", font_size=10.5)

    # 其他頁 header：顯示頁碼 PAGE
    normal_header = doc.sections[1].header
    if normal_header and normal_header.tables:
        tbl_other = normal_header.tables[0]
        cell_page = tbl_other.rows[1].cells[5]

        if cell_page.paragraphs:
            p_page = cell_page.paragraphs[0]
        else:
            p_page = cell_page.add_paragraph()

        clear_paragraph(p_page)

        add_simple_field(p_page, "PAGE", font_size=10.5)

# Extract plain text and "first seen" color from a cell JSON (your ProseMirror-ish structure).
def _extract_text_and_color_from_cell_json(cell_json, COLOR_DICT):
    txt_parts = []
    first_color = None
    for blk in cell_json.get("content", []):
        if blk.get("type") == "paragraph":
            for t in blk.get("content", []) or []:
                if t.get("type") == "text":
                    txt_parts.append(t.get("text", ""))
                    marks = t.get("marks")
                    if first_color is None and marks:
                        ckey = marks[0].get("attrs", {}).get("color")
                        if ckey in COLOR_DICT:
                            first_color = COLOR_DICT[ckey]
    return "".join(txt_parts).strip(), first_color

# Extract header texts for row0 (works if headers are in tableHeader/tableCell/customTableCell)
def _header_texts(row0_cells):
    headers = []
    for cell in row0_cells:
        if cell.get("type") not in ["customTableCell", "tableCell", "tableHeader"]:
            continue
        htxt = []
        for blk in cell.get("content", []):
            if blk.get("type") == "paragraph":
                for t in blk.get("content", []) or []:
                    if t.get("type") == "text":
                        htxt.append(t.get("text", ""))
        headers.append("".join(htxt).strip())
    return headers

# Find indices of the 5 special headers among ALL headers. Returns dict if all present, else None.
def _find_spec_header_indices(headers):
    pos = {}
    for name in SPEC_HDRS:
        try:
            pos[name] = headers.index(name)
        except ValueError:
            return None
    return pos

# Very tolerant float extractor (find first number in text like "12.3mm" or "±2.0")
_num_pat = re.compile(r"[-+]?\d+(?:\.\d+)?")
def _to_float_or_none(s):
    if s is None: return None
    m = _num_pat.search(str(s))
    return float(m.group(0)) if m else None

_calc_significant_digits = lambda s: 0 if "." not in s else len(s.split(".")[-1])
def _compose_value(set_txt, upper_txt, lower_txt):
    s = _to_float_or_none(set_txt)
    u = _to_float_or_none(upper_txt)
    l = _to_float_or_none(lower_txt)

    if u == None and l == None:
        return ""
    
    elif u == None:
        return f"≥{l}"

    elif l == None:
        return f"≤{u}"

    elif s == None:
        return f"{lower_txt}~{upper_txt}"
    
    elif set_txt == upper_txt and set_txt == lower_txt:
        return set_txt

    u_sd = _calc_significant_digits(upper_txt)
    l_sd = _calc_significant_digits(lower_txt)
    s_sd = _calc_significant_digits(set_txt)
    plus = round(u - s, max(s_sd, u_sd)) if s_sd + u_sd != 0 else int(u - s)
    minus = round(s - l, max(s_sd, l_sd)) if s_sd + l_sd != 0 else int(s - l)
    return f"-{minus}/+{plus}\n({lower_txt}~{upper_txt})"

def set_style_fonts(style, latin="Arial", east_asia="標楷體"):
    """
    Apply Latin (ascii/hAnsi/cs) and East Asian (eastAsia) fonts to a Word style.
    Works for paragraph and character styles.
    """
    # Set high-level python-docx font (affects ascii/hAnsi sometimes, but we set XML explicitly too)
    if hasattr(style, "font"):
        style.font.name = latin

    # Ensure <w:rPr> exists
    rPr = style._element.rPr
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        style._element.append(rPr)

    # Ensure <w:rFonts> exists
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)

    # Set Latin/complex + East Asian
    rFonts.set(qn("w:ascii"), latin)          # Latin
    rFonts.set(qn("w:hAnsi"), east_asia)          # Latin (alternate)
    rFonts.set(qn("w:cs"), latin)             # Complex script (safe to mirror Latin)
    rFonts.set(qn("w:eastAsia"), east_asia)   # East Asian (Chinese)

def apply_default_fonts(doc, latin="Arial", east_asia="標楷體"):
    # Normal
    set_style_fonts(doc.styles["Normal"], latin, east_asia)

    # Headings 1..9 (apply if present)
    for i in range(1, 10):
        name = f"Heading {i}"
        if name in doc.styles:
            set_style_fonts(doc.styles[name], latin, east_asia)

    # Common extra styles (optional; add any you use)
    for s in ["Title", "Subtitle", "Intense Quote", "Quote", "Strong", "Emphasis"]:
        if s in doc.styles:
            set_style_fonts(doc.styles[s], latin, east_asia)

def update_p(p, mapping):
    if len(p.runs) == 0:
        return
    
    full_text = "".join([pr.text for pr in p.runs])

    if mapping.get(full_text) == None:
        return
    
    if ("POINT" in full_text) or ("REASON" in full_text) or ("DOC_NAME" in full_text) or ("PROJECT" in full_text) or ("ITEM_TYPE" in full_text) or ("STYLE_NO" in full_text):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
    for r in p.runs:
        r.text = ""

    p.runs[0].text = str(mapping[full_text])
    set_run_node_text_font_style(p.runs[0])

def replace_in_footer(footer, title_mapping):
    for row in footer.tables[0].rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                update_p(p, title_mapping)

def replace_in_header(header, title_mapping):
    for row in header.tables[0].rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                update_p(p, title_mapping)

def set_repeat_table_header(row):
    """ set repeat table row on every new page"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row

# --- New Content Generation Helper Functions ---
def set_run_node_text_font_style(run_node):
    # 基本字型：英文 Arial，中文 標楷體
    run_node.font.name = "Arial"

    r = run_node._element  # CT_R
    # ⭐ 正確取得/建立 <w:rPr>
    rPr = r.get_or_add_rPr()

    # ⭐ 取得/建立 <w:rFonts>
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)

    # 設定字型：英文 + 中文
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rFonts.set(qn("w:eastAsia"), "標楷體")

def set_table_run_text(p_node, text, color = COLOR_DICT['#000']):
    run = p_node.add_run(text)
    run.font.color.rgb = color
    set_run_node_text_font_style(run)
    if color == COLOR_DICT["#0000ff"]:
        run.underline = True

    return run

def set_docx_table_cell_text(cell_node, text, color = COLOR_DICT["#000"], vCenter = False, center = False):
    p = cell_node.paragraphs[0]
    set_table_run_text(p, text, color)

    if vCenter:
        cell_node.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    return p

# ----- helper function for table merge processing and cell filling ----- #
def _extract_tiptap_table(json_table_data):
    """
    接受整個 doc 或 table node,
    回傳 table_node 與 rows_data(list of tableRow)
    """
    node_type = json_table_data.get("type")
    if node_type == "doc":
        # doc -> 第一個 table
        table_node = next(
            (n for n in json_table_data.get("content", []) if n.get("type") == "table"),
            None
        )
    elif node_type == "table":
        table_node = json_table_data
    else:
        table_node = None

    if not table_node:
        return None, []

    rows_data = [r for r in table_node.get("content", []) if r.get("type") == "tableRow"]
    return table_node, rows_data

def _compute_table_size(rows_data):
    """
    根據 TipTap 的 colspan / rowspan 計算邏輯上的 rowCount, colCount
    colCount = 每一 row 的 colspan 加總的最大值
    rowCount = rows_data 的長度
    """
    rowCount = len(rows_data)
    colCount = 0

    for row in rows_data:
        cells = [c for c in row.get("content", []) if c.get("type") in ["customTableCell", "tableCell", "tableHeader"]]
        col_span_sum = 0
        for cell in cells:
            attr = cell.get("attrs", {}) or {}
            col_span_sum += int(attr.get("colspan", 1) or 1)
        colCount = max(colCount, col_span_sum)

    return rowCount, colCount

def _build_span_grid(rows_data, rowCount, colCount):
    """
    建一個 rowCount x colCount 的矩陣 grid
    grid[r][c] = {
       "is_anchor": bool,
       "cell_data": cell_json or None,
       "rowspan": int,
       "colspan": int,
    }
    非 anchor 的被覆蓋格子也會填一格 {is_anchor: False, ...}
    """
    # 初始化為 None
    grid = [[None for _ in range(colCount)] for _ in range(rowCount)]

    for r_idx, row in enumerate(rows_data):
        cells = [c for c in row.get("content", []) if c.get("type") in ["customTableCell", "tableCell", "tableHeader"]]

        c_idx = 0  # 目前要填入的 column index
        for cell in cells:
            # 找到下一個空的 column
            while c_idx < colCount and grid[r_idx][c_idx] is not None:
                c_idx += 1
            if c_idx >= colCount:
                break  # 超出，理論上不會發生，防禦一下

            attr = cell.get("attrs", {}) or {}
            colspan = int(attr.get("colspan", 1) or 1)
            rowspan = int(attr.get("rowspan", 1) or 1)

            # 錨點
            grid[r_idx][c_idx] = {
                "is_anchor": True,
                "cell_data": cell,
                "rowspan": rowspan,
                "colspan": colspan,
            }

            # 填入被覆蓋的格子（非錨點）
            for dr in range(rowspan):
                for dc in range(colspan):
                    rr = r_idx + dr
                    cc = c_idx + dc
                    if rr == r_idx and cc == c_idx:
                        continue  # 錨點已填
                    if rr < rowCount and cc < colCount:
                        grid[rr][cc] = {
                            "is_anchor": False,
                            "cell_data": None,
                            "rowspan": 0,
                            "colspan": 0,
                        }

            c_idx += colspan

    return grid

def _fill_docx_cell_from_tiptap(docx_cell, cell_data, COLOR_DICT, center = True):
    """
    把 TipTap 的 cell_data 內容寫進 python-docx 的 cell
    保留你原本的 dropdown / paragraph / image 邏輯
    """
    attr = cell_data.get("attrs", {}) or {}
    dropdownValue = attr.get("dropdownValue")

    # 清掉 cell 內原本的空段落
    docx_cell.text = ""
    docx_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if dropdownValue:
        p = set_docx_table_cell_text(docx_cell, dropdownValue, color=COLOR_DICT.get(attr.get("dropdownColor")), vCenter=True, center=center)
        p.paragraph_format.keep_together = True
        return

    items = cell_data.get("content", []) or []
    first_para = True

    for item in items:
        if item.get("type") != "paragraph":
            continue

        # 第一個 paragraph 用現有的，之後用 add_paragraph
        p = docx_cell.paragraphs[0] if first_para and docx_cell.paragraphs else docx_cell.add_paragraph()
        first_para = False
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for text_item in item.get("content", []) or []:
            if text_item.get("type") == "text":
                marks = text_item.get("marks")
                run = p.add_run(text_item.get("text", ""))
                set_run_node_text_font_style(run)
                if marks:
                    # 假設只有一種 color mark
                    color_key = marks[0].get("attrs", {}).get("color")
                    if color_key in COLOR_DICT:
                        run.font.color.rgb = COLOR_DICT[color_key]
                        if COLOR_DICT[color_key] == RGBColor(0, 0, 255):
                            run.underline = True

            elif text_item.get("type") == "image":
                src = text_item["attrs"]["src"].split("/", 3)[-1]
                run = p.add_run()
                # 以目前 cell 寬度大概估一下圖片寬度
                cell_width_cm = docx_cell.width.cm if hasattr(docx_cell, "width") else 4
                run.add_picture(src, width=Cm(max(cell_width_cm - 0.5, 0.5)))

def create_docx_table(cell, json_table_data):
    rows_data = [row for row in json_table_data.get("content", []) if row.get("type") == "tableRow"]
    if not rows_data:
        return

    # Build list of header cells JSON (row 0)
    row0_cells_json = [c for c in rows_data[0].get("content", []) if c.get("type") in ["customTableCell", "tableCell", "tableHeader"]]
    if not row0_cells_json:
        return

    # Get header texts (full list, including extra headers like 項次/槽體/說明 etc.)
    headers = _header_texts(row0_cells_json)
    # Find positions of the 5 special headers; if not all present -> fall back to original rendering
    pos = _find_spec_header_indices(headers)

    # If NOT a spec table, render as-is (your original behavior)
    if pos is None:
        # 1) 抓 table / rows
        table_node, rows_data = _extract_tiptap_table(json_table_data)
        if not rows_data:
            return

        # 2) 計算 rowCount / colCount
        rowCount, colCount = _compute_table_size(rows_data)

        # 3) 建立 span grid
        grid = _build_span_grid(rows_data, rowCount, colCount)

        # 4) 建立 docx table（一次給足 rows / cols）
        table = cell.add_table(rows=rowCount, cols=colCount)
        table.style = "Table Grid"
        # 如果你有想控制寬度，可以在外面處理，這裡先略過
        # table.width = cell.width.cm - Cm(1)

        # 設定 header repeat（第一 row 視為 header）
        if rowCount > 0:
            set_repeat_table_header(table.rows[0])

        # 5) 先填內容，同時記錄 spanMap
        spanMap = {}  # (r,c) -> (rowspan, colspan)

        for r in range(rowCount):
            for c in range(colCount):
                info = grid[r][c]
                if not info or not info.get("is_anchor"):
                    continue  # 非錨點 cell，不用填內容，只等之後 merge

                cell_data = info["cell_data"]
                rowspan = int(info.get("rowspan", 1) or 1)
                colspan = int(info.get("colspan", 1) or 1)

                docx_cell = table.rows[r].cells[c]

                # 寫入實際內容
                _fill_docx_cell_from_tiptap(docx_cell, cell_data, COLOR_DICT)

                # 有合併才記錄
                if rowspan > 1 or colspan > 1:
                    spanMap[(r, c)] = (rowspan, colspan)

        # 6) 全部內容填完，再進行 merge
        for (r, c), (rowspan, colspan) in spanMap.items():
            first = table.cell(r, c)
            last = table.cell(r + rowspan - 1, c + colspan - 1)
            first.merge(last)

        return

    # ---------- SPEC TABLE TRANSFORMATION ----------
    # Determine insertion position = the smallest index among the five headers
    insert_at = min(pos.values())
    # The set of indices to remove (the five spec headers)
    spec_idx_set = set(pos[h] for h in SPEC_HDRS)

    # Build new header sequence: copy original headers left->right, but when we hit insert_at, insert ["規格值","操作值"] once and skip all five spec columns.
    new_header_titles = []
    col_map = []  # For non-spec columns: holds original index; for inserted pair: ('SPECPAIR', None)
    for index, header in enumerate(headers):
        # Replace SPEC column to calculated results
        if index == insert_at:
            col_map.append(("SPECPAIR", None, None))
            new_header_titles += ["設定值", "操作值", "規格值"]

        # Skip SPEC column expect first replaced one
        elif index in spec_idx_set:
            continue

        # Keep origin headers, if not SPEC column
        else:
            col_map.append(("COPY", index, len(new_header_titles)))
            new_header_titles.append(header)
    
    # If insert_at > len(headers) (theoretically impossible), guard
    if not new_header_titles:
        # fallback to original render if something odd happened
        return

    # Create the new docx table with computed column count
    table = cell.add_table(rows=0, cols=len(new_header_titles))
    table.width = cell.width.cm - Cm(1)
    table.style = 'Table Grid'

    # Write new header row
    hdr = table.add_row()
    set_repeat_table_header(hdr)
    for j, title in enumerate(new_header_titles):
        hdr.cells[j].text = title
        set_run_node_text_font_style(hdr.cells[j].paragraphs[0].runs[0])
        hdr.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr.cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Prepare a reverse map for quick lookup
    idx_reg_upper = pos["規格上限(OOS+)"]
    idx_reg_lower = pos["規格下限(OOS-)"]
    idx_opr_upper = pos["操作上限(OOC+)"]
    idx_opr_lower = pos["操作下限(OOC-)"]
    idx_set = pos["設定值"]

    NumColIndex = None if "項次" not in headers else headers.index('項次')
    previousNum = None if NumColIndex == None else 1

    SlotColIndex = None if "槽體" not in headers else headers.index('槽體')
    lastSlotName = None
    lastSlotRow = None if SlotColIndex == None else 1
    last_col_merges = []
    # print(f"rows_data: {rows_data}")

    # Now write each data row (skip row0 because it was header)
    for rowIndex, row_data in enumerate(rows_data):
        if rowIndex == 0:
            continue

        row = table.add_row()
        row.allow_row_break_across_pages = False
        row_cells_json = [c for c in row_data.get("content", []) if c.get("type") in ["customTableCell", "tableCell", "tableHeader"]]

        # Pre-extract text & color for all original columns (so we can reference easily)
        texts, colors = [], []
        for orig_idx, cjson in enumerate(row_cells_json):
            t, col = _extract_text_and_color_from_cell_json(cjson, COLOR_DICT)
            texts.append(t)
            colors.append(col)

        SlotName = texts[SlotColIndex]

        # Compose spec/oper values + color flags
        spec_val = _compose_value(texts[idx_set], texts[idx_reg_upper], texts[idx_reg_lower])
        oper_val = _compose_value(texts[idx_set], texts[idx_opr_upper], texts[idx_opr_lower])
        spec_is_red = COLOR_DICT['blue'] if (colors[idx_reg_upper] == COLOR_BLUE) or (colors[idx_reg_lower] == COLOR_BLUE) else COLOR_DICT['#000']
        oper_is_red = COLOR_DICT['blue'] if (colors[idx_opr_upper] == COLOR_BLUE) or (colors[idx_opr_lower] == COLOR_BLUE) else COLOR_DICT['#000']

        # Fill the new row respecting col_map
        for colIndex, (tag, dataIndex, dstIndex) in enumerate(col_map):
            # [NEW 2] 檢查是否為最後一欄，並讀取 rowspan
            # 只有當前欄位是最後一欄時才處理
            if colIndex == len(col_map) - 1:
                # 確保 dataIndex 在範圍內 (避免前端資料結構異常導致 crash)
                if dataIndex < len(row_cells_json):
                    curr_json = row_cells_json[dataIndex]
                    attrs = curr_json.get("attrs", {})
                    # 取得 rowspan，預設為 1
                    r_span = int(attrs.get("rowspan", 1) or 1)
                    
                    if r_span > 1:
                        # 記錄合併資訊: (開始列, 結束列, 目標欄位索引)
                        # rowIndex 對應 table.rows 的 index (因為 table 有 header 且 rows_data 有 header，剛好對齊)
                        start_r = rowIndex
                        end_r = rowIndex + r_span - 1
                        last_col_merges.append((start_r, end_r, dstIndex))
                        
            if tag == "SPECPAIR":
                p = set_docx_table_cell_text(row.cells[colIndex + 0], texts[idx_set], color = colors[idx_set], vCenter = True, center = True)
                p.paragraph_format.keep_together = True
                p = set_docx_table_cell_text(row.cells[colIndex + 1], oper_val, color = oper_is_red, vCenter = True, center = True)
                p.paragraph_format.keep_together = True
                p = set_docx_table_cell_text(row.cells[colIndex + 2], spec_val, color = spec_is_red, vCenter = True, center = True)
                p.paragraph_format.keep_together = True

            elif NumColIndex != None and colIndex == NumColIndex:
                if SlotName != lastSlotName:
                    p = set_docx_table_cell_text(row.cells[dstIndex], str(previousNum), color = colors[dataIndex], vCenter = True, center = True)
                    p.paragraph_format.keep_together = True
                    previousNum += 1

                    if lastSlotName != None and lastSlotRow != rowIndex - 1:
                        A = table.cell(lastSlotRow, dstIndex)
                        B = table.cell(rowIndex - 1, dstIndex)
                        A.merge(B)

            elif SlotColIndex != None and colIndex == SlotColIndex:
                if SlotName != lastSlotName:
                    p = set_docx_table_cell_text(row.cells[dstIndex], texts[dataIndex], color = colors[dataIndex], vCenter = True, center = True)
                    p.paragraph_format.keep_together = True

                    if lastSlotName != None and lastSlotRow != rowIndex - 1:
                        A = table.cell(lastSlotRow, dstIndex)
                        B = table.cell(rowIndex - 1, dstIndex)
                        A.merge(B)

                    lastSlotRow = rowIndex

            else:
                # 一般欄位寫入 (包含最後一欄的文字寫入)
                if dataIndex < len(texts):
                    p = set_docx_table_cell_text(row.cells[dstIndex], texts[dataIndex], color = colors[dataIndex], vCenter = True, center = True)
                    p.paragraph_format.keep_together = True

        lastSlotName = SlotName

    if lastSlotName != None and rowIndex != lastSlotRow:
        if NumColIndex != None:
            A = table.cell(lastSlotRow, NumColIndex)
            B = table.cell(rowIndex, NumColIndex)
            A.merge(B)
        
        A = table.cell(lastSlotRow, SlotColIndex)
        B = table.cell(rowIndex, SlotColIndex)
        A.merge(B)

    # [NEW 3] 執行最後一欄的合併
    # 這一步放在最後，避免干擾前面的文字寫入
    for start_r, end_r, col_idx in last_col_merges:
        try:
            # 防呆：確保結束列沒有超出表格範圍
            if end_r < len(table.rows):
                cell_start = table.cell(start_r, col_idx)
                cell_end = table.cell(end_r, col_idx)
                cell_start.merge(cell_end)
        except Exception as e:
            print(f"Error merging last column cells ({start_r}-{end_r}): {e}")

    return

def prevent_table_break(table):
    """
    設定表格盡量不跨頁 (Keep with next)。
    若表格長度超過一頁，Word 仍會強制換頁 (符合需求)。
    """
    # 1. 設定每一列 (Row) 內部的內容不要斷開 (Allow row to break across pages = False)
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        # 設定 <w:cantSplit>
        cantSplit = trPr.find(qn('w:cantSplit'))
        if cantSplit is None:
            cantSplit = OxmlElement('w:cantSplit')
            # 將 cantSplit 設為 on，表示「禁止」列內容跨頁
            cantSplit.set(qn('w:val'), 'on') 
            trPr.append(cantSplit)

    # 2. 設定「列與列之間」黏在一起 (Keep with next)
    # 邏輯：除了「最後一列」之外，每一列裡面的每一個段落都要設定 keep_with_next = True
    # 這樣 Row 1 會黏 Row 2, Row 2 黏 Row 3... 直到表格結束
    for row in table.rows[:-1]: # 排除最後一列
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True

def create_parameter_table(cell, condition_content, parameter_content, info):
    if (condition_content == None or len(condition_content) == 0) and (parameter_content == None or len(parameter_content) == 0):
        return
    
    # print(f"condition_content: {condition_content}")
    # print(f"prarmeter_content: {parameter_content}")
    
    programCode = '、'.join([p["programCode"] for p in info['programs']])
    machines_name = '、'.join([m['name'] for m in info['machines']])
    info_N_rows = 2 if info["step_type"] == 2 else 3

    table = None
    if condition_content != None and len(condition_content) != 0:
        tableNode = condition_content
        rows_data = [row for row in tableNode['content'][0]['content'] if row.get("type") == "tableRow"]
        cols = len(rows_data[0]['content'])

        table = cell.add_table(rows = info_N_rows + len(rows_data) + 1, cols = cols)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False

        total_width = cell.width - Cm(0.2)

        first_col_width = Inches(1)
        other_col_width = (total_width - first_col_width) // (cols - 1)
        first_col_width = int(total_width - (other_col_width * (cols - 1)))

        for row in table.rows:
            row.allow_row_break_across_pages = False

        for c in range(len(table.columns)):
            table.columns[c].width = first_col_width if c == 0 else other_col_width
            for table_cell in table.columns[c].cells:
                table_cell.width = first_col_width if c == 0 else other_col_width

        set_docx_table_cell_text(table.cell(0, 0), "程式代碼", center = True)
        table.cell(0, 1).merge(table.cell(0, cols - 1))
        set_docx_table_cell_text(table.cell(0, 1), programCode)
        set_docx_table_cell_text(table.cell(1, 0), "機台名稱", center = True)
        table.cell(1, 1).merge(table.cell(1, cols - 1))
        set_docx_table_cell_text(table.cell(1, 1), machines_name)
        set_docx_table_cell_text(table.cell(2, 0), "條件參數", center = True)
        table.cell(2, 0).merge(table.cell(2, cols - 1))

        for rowIndex, row_data in enumerate(rows_data):
            for colIndex, col_data in enumerate(row_data['content']):
                _fill_docx_cell_from_tiptap(table.cell(rowIndex + 3, colIndex), col_data, COLOR_DICT)

        prevent_table_break(table)

    write_info = (table == None)
    if parameter_content != None and len(parameter_content) != 0:
        tableNode = parameter_content
        rows_data = [row for row in tableNode['content'][0]['content'] if row.get("type") == "tableRow"]
        cols = len(rows_data[0]['content']) - 2

        if condition_content != None and len(condition_content) != 0:
            p_to_remove = cell.paragraphs[-1]._element  
            p_to_remove.getparent().remove(p_to_remove)
        
        table = cell.add_table(rows = info_N_rows + len(rows_data) + 1 if write_info else len(rows_data) + 1, cols = cols)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False

        total_width = cell.width - Cm(0.2)

        first_col_width = Inches(1)
        other_col_width = (total_width - first_col_width) // (cols - 1)
        first_col_width = int(total_width - (other_col_width * (cols - 1)))

        for row in table.rows:
            row.allow_row_break_across_pages = False

        for c in range(len(table.columns)):
            table.columns[c].width = first_col_width if c == 0 else other_col_width
            for table_cell in table.columns[c].cells:
                table_cell.width = first_col_width if c == 0 else other_col_width

        if write_info:
            set_docx_table_cell_text(table.cell(0, 0), "程式代碼", center = True)
            table.cell(0, 1).merge(table.cell(0, cols - 1))
            set_docx_table_cell_text(table.cell(0, 1), programCode)
            set_docx_table_cell_text(table.cell(1, 0), "機台名稱", center = True)
            table.cell(1, 1).merge(table.cell(1, cols - 1))
            set_docx_table_cell_text(table.cell(1, 1), machines_name)
            if info["step_type"] == 5:
                set_docx_table_cell_text(table.cell(2, 0), "流程順序", center = True)
                table.cell(2, 1).merge(table.cell(2, cols - 1))
                set_docx_table_cell_text(table.cell(2, 1), "、".join([f"第{order}次" for order in info["processOrder"]]))

        if info["step_type"] == 2:
            set_docx_table_cell_text(table.cell(2 if write_info else 0, 0), "製造參數", center = True)
            table.cell(2 if write_info else 0, 0).merge(table.cell(2 if write_info else 0, cols - 1))
        else:
            set_docx_table_cell_text(table.cell(3 if write_info else 0, 0), "製造參數", center = True)
            table.cell(3 if write_info else 0, 0).merge(table.cell(3 if write_info else 0, cols - 1))

        headers = _header_texts(rows_data[0]['content'])
        pos = _find_spec_header_indices(headers)

         # ---------- SPEC TABLE TRANSFORMATION ----------
        # Determine insertion position = the smallest index among the five headers
        insert_at = min(pos.values())
        # The set of indices to remove (the five spec headers)
        spec_idx_set = set(pos[h] for h in SPEC_HDRS)

        # Build new header sequence: copy original headers left->right, but when we hit insert_at, insert ["規格值","操作值"] once and skip all five spec columns.
        new_header_titles = []
        col_map = []  # For non-spec columns: holds original index; for inserted pair: ('SPECPAIR', None)
        for index, header in enumerate(headers):
            # Replace SPEC column to calculated results
            if index == insert_at:
                col_map.append(("SPECPAIR", None, None))
                new_header_titles += ["設定值", "操作值", "規格值"]

            # Skip SPEC column expect first replaced one
            elif index in spec_idx_set:
                continue

            # Keep origin headers, if not SPEC column
            else:
                col_map.append(("COPY", index, len(new_header_titles)))
                new_header_titles.append(header)
        
        # Write new header row
        for j, title in enumerate(new_header_titles):
            table.cell(info_N_rows + 1 if write_info else 1, j).text = title
            set_run_node_text_font_style(table.cell(info_N_rows + 1 if write_info else 1, j).paragraphs[0].runs[0])
            table.cell(info_N_rows + 1 if write_info else 1, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(info_N_rows + 1 if write_info else 1, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Prepare a reverse map for quick lookup
        idx_reg_upper = pos["規格上限(OOS+)"]
        idx_reg_lower = pos["規格下限(OOS-)"]
        idx_opr_upper = pos["操作上限(OOC+)"]
        idx_opr_lower = pos["操作下限(OOC-)"]
        idx_set = pos["設定值"]

        NumColIndex = None if "項次" not in headers else headers.index('項次')
        previousNum = None if NumColIndex == None else 1

        ExplainColIndex = None if "說明" not in headers else headers.index('說明')

        SlotColIndex = None if "槽體" not in headers else headers.index('槽體')
        lastSlotName = None
        lastSlotRow = None if SlotColIndex == None else 2
        lastSlotRow = lastSlotRow + info_N_rows if write_info else lastSlotRow

        for rowIndex, row_data in enumerate(rows_data):
            if rowIndex == 0:
                continue

            row = table.rows[rowIndex + info_N_rows + 1 if write_info else rowIndex + 1]
            row.allow_row_break_across_pages = False
            row_cells_json = [c for c in row_data.get("content", []) if c.get("type") in ["customTableCell", "tableCell", "tableHeader"]]

            # Pre-extract text & color for all original columns (so we can reference easily)
            texts, colors = [], []
            for orig_idx, cjson in enumerate(row_cells_json):
                t, col = _extract_text_and_color_from_cell_json(cjson, COLOR_DICT)
                texts.append(t)
                colors.append(col)

            SlotName = texts[SlotColIndex]

            # Compose spec/oper values + color flags
            spec_val = _compose_value(texts[idx_set], texts[idx_reg_upper], texts[idx_reg_lower])
            oper_val = _compose_value(texts[idx_set], texts[idx_opr_upper], texts[idx_opr_lower])
            spec_is_red = COLOR_DICT['blue'] if (colors[idx_reg_upper] == COLOR_BLUE) or (colors[idx_reg_lower] == COLOR_BLUE) else COLOR_DICT['#000']
            oper_is_red = COLOR_DICT['blue'] if (colors[idx_opr_upper] == COLOR_BLUE) or (colors[idx_opr_lower] == COLOR_BLUE) else COLOR_DICT['#000']

            rowIndex = rowIndex + info_N_rows + 1 if write_info else rowIndex + 1

            # Fill the new row respecting col_map
            for colIndex, (tag, dataIndex, dstIndex) in enumerate(col_map):
                if tag == "SPECPAIR":
                    p = set_docx_table_cell_text(row.cells[colIndex + 0], texts[idx_set], color = colors[idx_set], vCenter = True, center = True)
                    p.paragraph_format.keep_together = True
                    p = set_docx_table_cell_text(row.cells[colIndex + 1], oper_val, color = oper_is_red, vCenter = True, center = True)
                    p.paragraph_format.keep_together = True
                    p = set_docx_table_cell_text(row.cells[colIndex + 2], spec_val, color = spec_is_red, vCenter = True, center = True)
                    p.paragraph_format.keep_together = True

                elif NumColIndex != None and dataIndex == NumColIndex:
                    if SlotName != lastSlotName:
                        p = set_docx_table_cell_text(row.cells[dstIndex], str(previousNum), color = colors[dataIndex], vCenter = True, center = True)
                        p.paragraph_format.keep_together = True
                        previousNum += 1

                        if lastSlotName != None and lastSlotRow != rowIndex - 1:
                            A = table.cell(lastSlotRow, dstIndex)
                            B = table.cell(rowIndex - 1, dstIndex)
                            A.merge(B)

                elif SlotColIndex != None and dataIndex == SlotColIndex:
                    if SlotName != lastSlotName:
                        p = set_docx_table_cell_text(row.cells[dstIndex], texts[dataIndex], color = colors[dataIndex], vCenter = True, center = True)
                        p.paragraph_format.keep_together = True

                        if lastSlotName != None and lastSlotRow != rowIndex - 1:
                            A = table.cell(lastSlotRow, dstIndex)
                            B = table.cell(rowIndex - 1, dstIndex)
                            A.merge(B)

                        lastSlotRow = rowIndex

                elif ExplainColIndex != None and dataIndex == ExplainColIndex:
                    p = _fill_docx_cell_from_tiptap(row.cells[dstIndex], row_cells_json[dataIndex], COLOR_DICT = COLOR_DICT, center = False)

                else:
                    p = set_docx_table_cell_text(row.cells[dstIndex], texts[dataIndex], color = colors[dataIndex], vCenter = True, center = True)
                    p.paragraph_format.keep_together = True

            lastSlotName = SlotName

        if lastSlotName != None and rowIndex != lastSlotRow:
            if NumColIndex != None:
                A = table.cell(lastSlotRow, NumColIndex)
                B = table.cell(rowIndex, NumColIndex)
                A.merge(B)
            
            A = table.cell(lastSlotRow, SlotColIndex)
            B = table.cell(rowIndex, SlotColIndex)
            A.merge(B)
            
        prevent_table_break(table)

def set_cell_border(cell, **kwargs):
    """
    設定儲存格邊框的通用函式
    用法:
    set_cell_border(
        cell, 
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"}, 
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"val": "nil"}, # start 即 left (左邊框)
        end={"val": "nil"},   # end 即 right (右邊框)
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # 檢查 tcBorders 是否存在，不存在則建立
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # 處理傳入的邊框設定
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV', 'start', 'end'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # 如果該邊框設定已存在，先移除舊的
            element = tcBorders.find(qn(tag))
            if element is not None:
                tcBorders.remove(element)

            # 建立新的邊框節點
            border = OxmlElement(tag)
            
            # 設定屬性 (sz=大小, val=樣式, color=顏色, space=間距)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    border.set(qn('w:{}'.format(key)), str(edge_data[key]))
            
            tcBorders.append(border)

def create_process_table(cell, content_data):
    # print(f'content_data: {content_data}')
    if content_data.get('jsonContent') == None or content_data['jsonContent'].get('content') == None or content_data['jsonContent']['content'][0]:
        p_attr = cell.add_paragraph("NA")
        p_attr.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_attr.runs[0].font.size = Pt(12)
        return 
    tableNode = content_data['jsonContent']['content'][0]
    rows_data = [row for row in tableNode['content'] if row.get("type") == "tableRow"]
    if not rows_data:
        return []
    table = cell.add_table(rows=len(rows_data), cols=len(rows_data[0]['content']))
    table.width = cell.width.cm - Cm(1)
    table.style = 'Table Grid'

    for r, row_data in enumerate(rows_data):
        for c, cell_data in enumerate(row_data['content']):
            _fill_docx_cell_from_tiptap(table.rows[r].cells[c], cell_data, COLOR_DICT)

            if cell_data['content'][0].get('content') == None:
                set_cell_border(table.cell(r, c), left = {"val": "nil"}, top = {"val": "nil"}, right = {"val": "nil"}, bottom = {"val": "nil"})

    # p_to_remove = cell.paragraphs[-1]._element  
    # p_to_remove.getparent().remove(p_to_remove)

    # table2 = cell.add_table(rows=len(rows_data), cols=len(rows_data[0]['content']) - 1)
    # table2.width = cell.width.cm - Cm(1)
    # table2.style = 'Table Grid'

    # for r, row_data in enumerate(rows_data):
    #     for c, cell_data in enumerate(row_data['content']):
    #         if c == 9:
    #             continue
    #         _fill_docx_cell_from_tiptap(table2.rows[r].cells[c], cell_data, COLOR_DICT)

    #         if cell_data['content'][0].get('content') == None:
    #             set_cell_border(table2.cell(r, c), left = {"val": "nil"}, top = {"val": "nil"}, right = {"val": "nil"}, bottom = {"val": "nil"})

def parse_json_content(parent_object, json_data, indent = None, header = False, no = None, tier=1):
    """
    Recursively parses the JSON content (jsonHeader/jsonContent) and adds elements 
    (paragraphs, list items) to the parent object (docx.table._Cell).
    Applies tier-based indentation.
    """
    if json_data is None or json_data.get("content") is None:
        return

    base_indent = Cm(0.7) 

    for block in json_data["content"]:
        block_type = block.get("type")
        
        # Use the 'tier' from the content object, or the current recursion tier.
        content_tier = block.get("tier", tier) 

        if block_type == "paragraph" or block_type == "listItem":
            p = parent_object.add_paragraph()
            p.style = 'List Paragraph' if block_type == "listItem" else 'Normal'

            if header:
                set_table_run_text(p, f"{no} ")

            # Set text for title in content table
            for item in block.get("content", []):
                if item.get("type") == "text":
                    marks = item.get("marks", [])
                    color = COLOR_DICT["#000"]
                    if len(marks) > 0:
                        color_key = marks[0].get("attrs", {}).get("color")
                        if color_key in COLOR_DICT:
                            color = COLOR_DICT[marks[0].get("attrs", {"color": "#000"}).get("color")] if len(marks) > 0 else COLOR_DICT["#000"]
                    run = set_table_run_text(p, item.get("text", ""), color)
                    set_run_node_text_font_style(run)
                    run.font.size = Pt(12)


            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Apply tier-based indent. tier=1 means 1 level of content under step title.
            # Indent factor is tier - 1. (Tier 1 gets indent factor 0, Tier 2 gets 1, etc.)
            indent_factor = max(0, content_tier - 1)
            p.paragraph_format.left_indent = base_indent * indent_factor
        
        elif block_type == "table":
            # For tables nested inside content structure (less common but handled)
            create_docx_table(parent_object, block)
            
        # Recursive call for nested content
        if 'content' in block and block.get("type") not in ["table", "tableRow", "tableCell", "tableHeader"]:
            # Pass the current tier + 1 for deeper nesting
            parse_json_content(parent_object, block, content_tier + 1)


def createHeader(cell, step_content_list):
    """Adds header content (jsonHeader) for all items in the list to the cell."""
    for content_obj in step_content_list:
        no = f"{content_obj['step']}.{content_obj['tier']}" if content_obj['sub_no'] == 0 else f"{content_obj['step']}.{content_obj['tier']}.{content_obj['sub_no']}"
        tier = 2 if content_obj.get("sub_no", 0) == 0 else 3
        for item in content_obj.get("data", []):
            if item.get("jsonHeader"):
                # Headers are main headings, usually tier 1 (no base indent)
                parse_json_content(cell, item["jsonHeader"], header = True, no = no, tier = tier) 

def createContent(cell, step_content_list):
    """Adds main content (jsonContent) for all items in the list to the cell, applying tier-based indent."""
    for content_obj in step_content_list:
        tier = 3 if content_obj.get("sub_no", 0) == 0 else 4
        for item in content_obj.get("data", []):
            if item.get("jsonContent"):
                # Pass the content's tier to enable indentation
                parse_json_content(cell, item["jsonContent"], tier=tier)

def createPictures(cell, step_content_list):
    """Adds picture placeholders/links (files) for all items in the list to the cell."""
    width = cell.width.cm
    base_indent = Cm(0.7)
    for content_obj in step_content_list:
        tier = content_obj.get("tier", 1)
        indent_factor = max(0, tier - 1)
        for item in content_obj.get("data", []):
            if item.get("files"):
                for file_info in item["files"]:
                    # Cannot embed image, so use a placeholder text/link
                    # p = cell.add_paragraph(f"圖片連結: {file_info.get('url', 'N/A')}")
                    # p.runs[0].font.size = Pt(10)
                    # p.paragraph_format.left_indent = base_indent * indent_factor
                    if file_info.get("path_to_save") != None:
                        p = cell.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        src = file_info["path_to_save"].split("/", 1)[-1]
                        run = p.add_run()
                        run.add_picture(f'uploads/temp/{src}', width = Cm(width / 2 - 0.5))

def createTable(cell, step_content_list, info = {}):
    """Adds tables (jsonContent containing a table) for all items in the list to the cell."""
    for content_obj in step_content_list:
        for item in content_obj.get("data", []):
            if item.get("jsonContent"):
                # The JSON for table content is expected to be a single 'table' block inside 'content'
                table_blocks = [b for b in item["jsonContent"].get("content", []) if b.get("type") == "table"]
                for table_block in table_blocks:
                    # Table is created directly in the cell, as requested, to avoid left indent.
                    create_docx_table(cell, table_block)

        if info.get("step_type"):
            create_parameter_table(cell, content_obj.get("jsonConditionContent", []), content_obj.get("jsonParameterContent", []), info)

def draw_instruction_content(doc, data):
    attribute = data["attribute"][-1]
    contents = data["content"]
    references = data["reference"]

    documentType = attribute["documentType"]
    cell = doc.tables[1].rows[0].cells[0]
    
    # Ensure a paragraph exists to start or continue content in the cell
    if not cell.paragraphs or cell.paragraphs[0].text:
         cell.add_paragraph()
    cell.paragraphs[0].style = doc.styles["Normal"]

    # Get document chapters from definition
    for (stepIndex, itemInfo) in enumerate(DOCUMENT_STEP[DOCUMENT_TYPE(documentType).name].value):
        (step, stepInfo), = itemInfo.items()
        
        # 1. Set the step title
        if stepIndex > 0:
            cell.add_paragraph()
        p_title = cell.paragraphs[0] if stepIndex == 0 else cell.add_paragraph()
        p_title.text = f"{stepIndex + 1}.{step}"
        p_title.runs[0].font.size = Pt(12)
        set_run_node_text_font_style(p_title.runs[0])
        p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 2. Get current step content
        if stepInfo["parent"] == "attribute":
            stepContent = attribute.get(stepInfo["code"])
            if isinstance(stepContent, str):
                stepContent = "NA" if len(stepContent) == 0 else stepContent
                p_attr = cell.add_paragraph(stepContent)
                p_attr.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_attr.runs[0].font.size = Pt(12)
                set_run_node_text_font_style(p_attr.runs[0])
                p_attr.paragraph_format.left_indent = p_attr.runs[0].font.size * 2
        
        elif stepInfo["parent"] == "content":
            # Filter content items for the current step_type
            step_content_list = [content for content in contents if content["step_type"] == stepInfo["code"]]

            if not step_content_list:
                stepContent = "NA" if len(step_content_list) == 0 else stepContent
                p_attr = cell.add_paragraph(stepContent)
                p_attr.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_attr.runs[0].font.size = Pt(12)
                set_run_node_text_font_style(p_attr.runs[0])
                p_attr.paragraph_format.left_indent = p_attr.runs[0].font.size * 2
                continue

            # Process all data items across all matching step content objects
            for content_obj in step_content_list:
                if content_obj["step_type"] == 2 or content_obj["step_type"] == 5:
                    machines, processOrder = [], None
                    if content_obj["step_type"] == 2:
                        machines_name = attribute["attribute"]["inputMachines"].split(",")
                        machines_code = attribute["attribute"]["machines"]
                        machines = [{"code": code, "name": name} for code, name in zip(machines_code, machines_name)]
                    elif content_obj["step_type"] == 5:
                        machines_name = content_obj['metadata']['machines_name']
                        machines_code = content_obj['metadata']['machines']
                        machines = [{"code": code, "name": name} for code, name in zip(machines_code, machines_name)]
                        processOrder = content_obj['metadata']['processOrder']
                    
                    createTable(cell, [content_obj], {"step_type": content_obj["step_type"], "programs": content_obj["metadata"]['programs'], "machines": machines, "processOrder": processOrder})

                if content_obj["step_type"] == 0:
                    item_for_helper = {"data": content_obj.get("data", []), "step": stepIndex + 1, "tier": content_obj.get("tier", 1), "sub_no": 0}
                    createHeader(cell, [item_for_helper]) 
                    create_process_table(cell, content_obj['data'][0])
                    continue
                   
                if content_obj.get('step_type') == 1 and content_obj.get('tier') == 1:
                    stepContent = "3.1 生產管理條件"
                    p_attr = cell.add_paragraph(stepContent)
                    p_attr.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_attr.runs[0].font.size = Pt(12)
                    set_run_node_text_font_style(p_attr.runs[0])
                    p_attr.paragraph_format.left_indent = p_attr.runs[0].font.size * 2

                    tiptap_content = content_obj["data"]
                    content = "" if tiptap_content[0]['jsonHeader']['content'][0].get("content") == None else tiptap_content[0]['jsonHeader']['content'][0]["content"][0]['text']
                    if len(content) > 0:
                        p_attr = cell.add_paragraph(content)
                        p_attr.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p_attr.runs[0].font.size = Pt(12)
                        set_run_node_text_font_style(p_attr.runs[0])
                        p_attr.paragraph_format.left_indent = p_attr.runs[0].font.size * 4
                    item_for_helper = {"data": tiptap_content, "step": 3, "tier": content_obj.get("tier", 1), "sub_no": 0}
                    createTable(cell, [item_for_helper]) 
                    continue

                for index, item_data in enumerate(content_obj.get("data", [])):
                    option = item_data.get("option")
                    
                    # Wrap the single data item for consistent access by helpers
                    item_for_helper = {"data": [item_data], "step": stepIndex + 1, "tier": content_obj.get("tier", 1), "sub_no": index}
                    
                    if option == 2: # Header + Table
                        createHeader(cell, [item_for_helper]) 
                        createTable(cell, [item_for_helper]) 
                    elif option == 1: # Header + Content + Pictures
                        createHeader(cell, [item_for_helper]) 
                        createContent(cell, [item_for_helper]) 
                        createPictures(cell, [item_for_helper]) 
                    elif option == 0: # Header only
                        createHeader(cell, [item_for_helper])
            
        elif stepInfo["parent"] == "reference":
            # Filter reference items for the current referenceType
            stepContent = [reference for reference in references if reference["referenceType"] == stepInfo["code"]]
            
            if stepContent:
                # Add a reference list with indentation
                for index, ref in enumerate(stepContent):
                    # For reference documents, use a simple numbered format
                    ref_p = cell.add_paragraph(f"({index + 1}) {ref.get('referenceDocumentID', '')} - {ref.get('referenceDocumentName', '')}")
                    ref_p.runs[0].font.size = Pt(12)
                    ref_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    ref_p.paragraph_format.left_indent = ref_p.runs[0].font.size * 2
                    ref_p.paragraph_format.space_after = Pt(3)
                    set_run_node_text_font_style(ref_p.runs[0])

            else:
                stepContent = "NA" if len(stepContent) == 0 else stepContent
                p_attr = cell.add_paragraph(stepContent)
                p_attr.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_attr.runs[0].font.size = Pt(12)
                p_attr.paragraph_format.left_indent = p_attr.runs[0].font.size * 2
                set_run_node_text_font_style(p_attr.runs[0])

def generate_word_password_hash(password, spin_count=100000):
    """
    依照 Office Open XML (OOXML) 標準產生密碼的 salt 與 hash
    """
    # 1. 產生 16 bytes 的隨機鹽值
    salt = os.urandom(16)
    
    # 2. Word 規定密碼必須使用 UTF-16 LE 編碼
    pwd_bytes = password.encode('utf-16le')
    
    # 3. 初始雜湊：SHA512(salt + password)
    hash_calc = hashlib.sha512(salt + pwd_bytes).digest()
    
    # 4. 進行 Spin 迭代運算 (預設 10 萬次)
    for i in range(spin_count):
        # 迭代器必須是 4 bytes 的 Little-Endian 格式
        iterator = i.to_bytes(4, byteorder='little')
        hash_calc = hashlib.sha512(iterator + hash_calc).digest()
        
    # 5. 回傳 Base64 編碼後的字串
    salt_b64 = base64.b64encode(salt).decode('utf-8')
    hash_b64 = base64.b64encode(hash_calc).decode('utf-8')
    
    return salt_b64, hash_b64

def enable_docx_protection(doc, password):
    """
    直接修改 python-docx 的 document 物件，注入保護設定與密碼驗證參數。
    """
    settings_element = doc.settings.element

    # 取得符合 Word 標準的 salt 與 hash
    salt_value, hash_value = generate_word_password_hash(password)

    protection = OxmlElement('w:documentProtection')
    
    # 設定保護類型
    protection.set(qn('w:edit'), 'forms')  
    protection.set(qn('w:enforcement'), '1')
    
    # 設定加密參數
    protection.set(qn('w:cryptProviderType'), 'rsaAES')
    protection.set(qn('w:cryptAlgorithmClass'), 'hash')
    protection.set(qn('w:cryptAlgorithmType'), 'typeAny')
    protection.set(qn('w:cryptAlgorithmSid'), '14') # 14 代表 SHA-512
    protection.set(qn('w:cryptSpinCount'), '100000')
    
    # ★ 關鍵修正：必須同時寫入 salt 與 hash
    protection.set(qn('w:salt'), salt_value)
    protection.set(qn('w:hash'), hash_value)

    # 移除舊標籤
    existing = settings_element.find(qn('w:documentProtection'))
    if existing is not None:
        settings_element.remove(existing)

    # 寫入新標籤
    settings_element.append(protection)

def fill_from_template(template_path, out_path, data, title_mapping, info_mapping):
    doc = Document(template_path)
    apply_default_fonts(doc, latin="Arial", east_asia="標楷體")

    # 產生內容（影響頁數）
    draw_instruction_content(doc, data)

    # 先做原本的 header/footer 文字替換
    for index, section in enumerate(doc.sections):
        # 如果你的第一頁 header 也要做文字替換，可以另外處理：
        # replace_in_header(section.first_page_header, title_mapping)  # 視需求決定
        replace_in_header(section.header, title_mapping)
        replace_in_footer(section.footer, title_mapping)
        # print(f"index: {index}, section: {section.header}")

    for row in doc.tables[0].rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                update_p(p, info_mapping)

    # ★ 內容都產生完之後，再插入頁碼欄位
    setup_page_numbers(doc)

    enable_docx_protection(doc, "123456")

    doc.save(out_path)

_code_prefix_re = re.compile(r"^\s*\(([^)]+)\)\s*(.*)$")

def clean_process_name(raw: str) -> str:
    """
    PROCESS_NAME like "(L100-01)單面前處理" -> "單面前處理"
    If there is no "(...)" prefix, returns the original, trimmed.
    """
    if not raw:
        return ""
    m = _code_prefix_re.match(raw)
    return m.group(2).strip() if m else raw

def get_docx(outpath, data, template = "docx-template/example3.docx"):
    attribute = data["attribute"][-1]
    Doc_id = attribute["documentID"]
    Date = datetime.now().strftime("%Y/%m/%d")
    Version = f"{int(attribute['documentVersion']):.1f}"
    Title = "製造條件指示書" if attribute["documentType"] == 0 else "製造式樣書"
    Doc_name = attribute["documentName"]

    itemType = attribute["attribute"].get("itemType")
    styleNo = attribute["attribute"].get("styleNo")

    # 👇 新增 documentKey（每次下載的唯一 key）
    doc_key = attribute.get("documentKey", "")

    title_mapping = {
        "DOC_NO": Doc_id,
        "DATE": Date,
        "REV": Version,
        "PAGE": "1",
        "TITLE": Title,
        "DOC_NAME": Doc_name,
        "PROJECT": "",
        "ITEM_TYPE": "",
        "STYLE_NO": "",
        "DOC_KEY": doc_key,
    }
    if itemType is not None:
        specifications = [clean_process_name(s["name"]) for s in attribute["attribute"].get("specification")]
        title_mapping["PROJECT"] = "_".join(specifications)
        title_mapping["ITEM_TYPE"]=  itemType
        title_mapping["STYLE_NO"] = styleNo.split("-", 1)[-1]
    else:
        title_mapping["PROJECT"] = attribute["attribute"].get("applyProject", "")

    info_mapping = {
        "REV1": "", "DATE1": "", "REASON1": "", "POINT1": "", "DEPT1": "", "APPROVER1": "", "CONFIRMER1": "", "AUTHOR1": "",
        "REV2": "", "DATE2": "", "REASON2": "", "POINT2": "", "DEPT2": "", "APPROVER2": "", "CONFIRMER2": "", "AUTHOR2": "",
        "REV3": "", "DATE3": "", "REASON3": "", "POINT3": "", "DEPT3": "", "APPROVER3": "", "CONFIRMER3": "", "AUTHOR3": ""
    }

    # 若你要在內容 table 裡也顯示 key，可以這樣放：
    # info_mapping["DOC_KEY"] = doc_key

    # Place document attribute into Word
    for index, attribute in enumerate(data["attribute"]):
        info_mapping[f"REV{index + 1}"] = f'{attribute["documentVersion"]:.1f}'
        info_mapping[f"DATE{index + 1}"] = attribute.get("issueDate", datetime.now().strftime("%Y/%m/%d"))
        info_mapping[f"REASON{index + 1}"] = attribute["reviseReason"]
        info_mapping[f"POINT{index + 1}"] = attribute["revisePoint"]
        info_mapping[f"DEPT{index + 1}"] = attribute["department"]
        info_mapping[f"APPROVER{index + 1}"] = attribute["approver"]
        info_mapping[f"CONFIRMER{index + 1}"] = attribute["confirmer"]
        info_mapping[f"AUTHOR{index + 1}"] = attribute["author"]
        info_mapping[f"REASON{index + 1}"] = f'變更理由\n{attribute["reviseReason"]}'
        info_mapping[f"POINT{index + 1}"] = f'變更要點\n{attribute["revisePoint"]}'

    # Assuming 'example__.docx' exists in the execution environment
    fill_from_template(template, outpath, data, title_mapping, info_mapping) 

# ----------------- Example usage -----------------
if __name__ == "__main__":
    # The file path has been changed to match the payload.json fetched content
    # Note: Replace '_captures/docs/20251028-114010-2b096f8a/payload.json' 
    # with the actual file path if running locally. 
    # Since I have the content, I will use the dictionary directly for demonstration.
    # with open("_captures/docs/20251030-161207-516a0cfa/payload.json", 'r', encoding = "utf-8") as f:
    with open("_captures/docs/20251103-095207-27ee69e9/payload.json", 'r', encoding = "utf-8") as f:
        data = json.load(f)

    # template = r"./example__.docx"   # <- your example file (the “second” screenshot)
    # output   = r"./header_table_fix.docx"
    template = "docx-template/example3.docx"
    output = "docxTemp/temp.docx"

    attribute = data["attribute"][-1]
    Doc_id = attribute["documentID"]
    Date = datetime.now().strftime("%Y/%m/%d")
    Version = f"{int(attribute['documentVersion']):.1f}"
    Title = "製造條件指示書" if attribute["documentType"] == 0 else "製造式樣書"
    Doc_name = attribute["documentName"]

    # Put placeholders like [DOC_NO], [REV], [DATE], [TOTAL_PAGES] in the header cells of your template.
    # title_mapping = { "DOC_NO": "WMH250", "DATE": "2025/10/29", "REV": "1.0", "PAGE": "1", "TITLE": "製造條件指示書", "DOC_NAME": "", "PROJECT": "", "DOC_CODE": "FM-R-MF-AZ-052 Rev7.0"}
    title_mapping = { "DOC_NO": Doc_id, "DATE": Date, "REV": Version, "PAGE": "1", "TITLE": Title, "DOC_NAME": Doc_name, "PROJECT": "", "DOC_CODE": "FM-R-MF-AZ-052 Rev7.0"}
    info_mapping = {
        "REV1": "", "DATE1": "", "REASON1": "", "POINT1": "", "DEPT1": "", "APPROVER1": "", "CONFIRMER1": "", "AUTHOR1": "",
        "REV2": "", "DATE2": "", "REASON2": "", "POINT2": "", "DEPT2": "", "APPROVER2": "", "CONFIRMER2": "", "AUTHOR2": "",
        "REV3": "", "DATE3": "", "REASON3": "", "POINT3": "", "DEPT3": "", "APPROVER3": "", "CONFIRMER3": "", "AUTHOR3": ""
    }

    for index, attribute in enumerate(data["attribute"]):
        info_mapping[f"REV{index + 1}"] = f'{attribute["documentVersion"]:.1f}'
        info_mapping[f"DATE{index + 1}"] = attribute.get("issueDate", datetime.now().strftime("%Y/%m/%d"))
        info_mapping[f"REASON{index + 1}"] = attribute["reviseReason"]
        info_mapping[f"POINT{index + 1}"] = attribute["revisePoint"]
        info_mapping[f"DEPT{index + 1}"] = attribute["department"]
        info_mapping[f"APPROVER{index + 1}"] = attribute["approver"]
        info_mapping[f"CONFIRMER{index + 1}"] = attribute["confirmer"]
        info_mapping[f"AUTHOR{index + 1}"] = attribute["author"]
        info_mapping[f"REASON{index + 1}"] = f'變更理由\n{attribute["reviseReason"]}'
        info_mapping[f"POINT{index + 1}"] = f'變更要點\n{attribute["revisePoint"]}'

    # Assuming 'example__.docx' exists in the execution environment
    fill_from_template(template, output, data, title_mapping, info_mapping) 
    print("The script with the completed draw_instruction_content function is ready.")
    print("Please ensure your local environment has 'example__.docx' and python-docx installed to run the full script.")
    